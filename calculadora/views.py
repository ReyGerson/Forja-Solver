"""
Vistas principales de la aplicación Forja-Solver.
Incluye lógica para métodos numéricos (Punto Fijo, Trazador Cúbico), gestión de usuarios, historial, premium y generación de PDFs.
"""
from django.shortcuts import render, redirect
import json
from django.http import HttpResponse
from .forms import PuntoFijoForm, SplineInputForm, RegistroUsuarioForm, EditarPerfilForm
from .models import SplineHistory, PuntoFijoHistorial, SimplexHistorial, MetodoGraficoHistorial
from .utils import parse_points, natural_cubic_spline
import numexpr as ne
from reportlab.pdfgen import canvas
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from .user_profile import UserProfile
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required
import re
from typing import List, Tuple
import pandas as pd
import pandas as pd
import os
import time

@login_required
def punto_fijo_view(request):
    """
    Vista principal para el método de Punto Fijo.
    Permite calcular la raíz de una función usando un despeje g(x).
    Guarda el historial y muestra la gráfica solo a usuarios premium.
    """
    resultado = []
    solucion = None
    error = None
    comprobacion = None
    formula_funcion = ""
    formula_despeje = ""

    if request.method == 'POST':
        form = PuntoFijoForm(request.POST)
        if form.is_valid():
            fx_input = form.cleaned_data['funcion']
            gx_input = form.cleaned_data['despeje']
            x0 = form.cleaned_data['valor_inicial']
            decimales = form.cleaned_data['decimales']
            tolerancia = form.cleaned_data['tolerancia']

            formula_funcion = fx_input
            formula_despeje = gx_input

            iteraciones = []
            i = 0
            ea = 100
            x1 = None

            while True:
                i += 1
                try:
                    x1 = float(ne.evaluate(gx_input, local_dict={'x': x0}))
                    fxi = float(ne.evaluate(fx_input, local_dict={'x': x1}))

                    if i > 1:
                        ea = abs((x1 - x0) / x1) * 100
                    else:
                        ea = None

                    valor_x0 = round(x0, decimales)
                    valor_x1 = round(x1, decimales)
                    detalle = f"x{i} = {gx_input.replace('x', str(valor_x0))} = {valor_x1}"
                    if ea is not None:
                        error_detalle = f"ε{i} = |({valor_x1} - {valor_x0}) / {valor_x1}| × 100% = {round(ea, 2)}%"
                    else:
                        error_detalle = "ε1 = -"

                    iteraciones.append({
                        'iteracion': i,
                        'detalle': detalle,
                        'error_detalle': error_detalle,
                        'x': valor_x1,
                        'error': round(ea, 2) if ea is not None else '-',
                    })

                    if ea is not None and ea <= tolerancia:
                        break

                    x0 = x1

                except Exception as e:
                    form.add_error(None, f"Error al evaluar la función o el despeje: {e}")
                    iteraciones = []
                    x1 = None
                    break

            if iteraciones and x1 is not None:
                solucion = round(x1, decimales)
                error = round(ea, 5)
                try:
                    comprobacion = round(float(ne.evaluate(fx_input, local_dict={'x': solucion})), 8)
                except:
                    comprobacion = "Error al comprobar"

                resultado = iteraciones

                # Guardar historial
                iteraciones_texto = "\n".join(
                    f"Iter {r['iteracion']}: x={r['x']} | Error={r['error']}%" for r in resultado
                )

                PuntoFijoHistorial.objects.create(
                    user=request.user,
                    funcion=fx_input,
                    despeje=gx_input,
                    valor_inicial=form.cleaned_data['valor_inicial'],
                    tolerancia=tolerancia,
                    decimales=decimales,
                    solucion=solucion,
                    error=error,
                    comprobacion=str(comprobacion),
                    iteraciones=iteraciones_texto,
                    funcion_latex=request.POST.get('funcion_latex', ''),
                    despeje_latex=request.POST.get('despeje_latex', ''),
                )

    else:
        # Set initial values for the form so form.initial always has the keys
        form = PuntoFijoForm(initial={
            'valor_inicial': 1,
            'tolerancia': 0.001,
            'decimales': 6,
        })

    # --- Fix: ensure form.data always has the needed keys for template ---
    # Use try/except to avoid errors if form.data is not a dict (e.g., QueryDict)
    try:
        if 'valor_inicial' not in form.data or not form.data['valor_inicial']:
            form.data = form.data.copy()
            form.data['valor_inicial'] = form.initial.get('valor_inicial', 1)
        if 'tolerancia' not in form.data or not form.data['tolerancia']:
            form.data['tolerancia'] = form.initial.get('tolerancia', 0.001)
        if 'decimales' not in form.data or not form.data['decimales']:
            form.data['decimales'] = form.initial.get('decimales', 6)
    except Exception:
        pass
    # --- end fix ---

    # --- Definir valores seguros para los campos del formulario ---
    valor_inicial = form.data.get('valor_inicial') or form.initial.get('valor_inicial', 1)
    tolerancia = form.data.get('tolerancia') or form.initial.get('tolerancia', 0.001)
    decimales = form.data.get('decimales') or form.initial.get('decimales', 6)
    # --- end fix ---

    # --- Datos para gráfica Punto Fijo ---
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
    grafica_punto_fijo = None
    if resultado and is_premium:
        grafica_punto_fijo = {
            'x': [r['iteracion'] for r in resultado],
            'y': [r['x'] for r in resultado],
            'error': [r['error'] for r in resultado],
        }
    # --- end datos gráfica ---

    print("POST recibido:", request.POST)
    return render(request, 'punto_fijo/puntoFijo.html', {
        'form': form,
        'resultado': resultado,
        'solucion': solucion,
        'error': error,
        'comprobacion': comprobacion,
        'fx': formula_funcion,
        'gx': formula_despeje,
        'fx_latex': request.POST.get('funcion_latex', ''), 
        'gx_latex': request.POST.get('despeje_latex', ''),
        'is_premium': is_premium,
        'valor_inicial_val': valor_inicial,
        'tolerancia_val': tolerancia,
        'decimales_val': decimales,
        'grafica_punto_fijo': grafica_punto_fijo,
    })

@login_required
def spline_view(request):
    """
    Vista principal para el método de Trazador Cúbico.
    Permite interpolar un valor usando splines cúbicos naturales.
    Guarda el historial y muestra la gráfica solo a usuarios premium.
    """
    result = None
    if request.method == 'POST':
        form = SplineInputForm(request.POST)
        if form.is_valid():
            points_str = form.cleaned_data['points']
            x_val = form.cleaned_data['x_value']
            try:
                points = parse_points(points_str)
                result = natural_cubic_spline(points, x_val)

                # Guardar en base de datos
                SplineHistory.objects.create(
                    user=request.user,         
                    puntos=points_str,
                    x_valor=x_val,
                    resultado=result['resultado'],
                    razonamiento=result['razonamiento'],
                    polinomio_usado=result['polinomio_usado']
                )
            except Exception as e:
                form.add_error(None, f"Error en los datos: {e}")
    else:
        form = SplineInputForm()

    # --- Datos para gráfica Trazador Cúbico ---
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
    grafica_trazador = None
    if result and 'tabla_hm' in result and is_premium:
        try:
            # Extraer puntos para graficar la curva y los puntos originales
            puntos = []
            if hasattr(form, 'cleaned_data') and 'points' in form.cleaned_data:
                puntos = parse_points(form.cleaned_data['points'])
            elif hasattr(form, 'initial') and 'points' in form.initial:
                puntos = parse_points(form.initial['points'])
            x_vals = [p[0] for p in puntos]
            y_vals = [p[1] for p in puntos]
            # Si hay polinomios, graficar la curva
            polinomios = result.get('polinomios', [])
            grafica_trazador = {
                'x': x_vals,
                'y': y_vals,
                'polinomios': polinomios,
                'x_query': result.get('x_query'),
                'y_query': result.get('resultado'),
            }
        except Exception:
            grafica_trazador = None
    # --- end datos gráfica ---

    return render(request, 'trazador_cubico/trazadorCubico.html', {
        'form': form,
        'result': result,
        'is_premium': is_premium,
        'grafica_trazador': grafica_trazador,
    })

@login_required
def historial_view(request):
    """
    Muestra el historial de cálculos de Trazador Cúbico del usuario.
    """
    historial = SplineHistory.objects.filter(user=request.user).order_by('-fecha_creacion')
    return render(request, 'trazador_cubico/historialTrazador.html', {'historial': historial})

@login_required
def historial_punto_fijo(request):
    """
    Muestra el historial de cálculos de Punto Fijo del usuario.
    """
    historial = PuntoFijoHistorial.objects.filter(user=request.user).order_by('-fecha')
    return render(request, 'punto_fijo/historial_punto_fijo.html', {'historial': historial})

@login_required
def punto_fijo_pdf(request, id):
    """
    Genera un PDF con el historial de un cálculo de Punto Fijo.
    """
    obj = PuntoFijoHistorial.objects.get(id=id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="punto_fijo_{obj.id}.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)
    y = 800

    p.drawString(50, y, f"PUNTO FIJO - Historial ID {obj.id}")
    y -= 30
    p.drawString(50, y, f"F(x): {obj.funcion}")
    y -= 20
    p.drawString(50, y, f"g(x): {obj.despeje}")
    y -= 20
    p.drawString(50, y, f"x₀: {obj.valor_inicial}, Tolerancia: {obj.tolerancia}, Decimales: {obj.decimales}")
    y -= 20
    p.drawString(50, y, f"Solución: x = {obj.solucion}, Error = {obj.error}%")
    y -= 20
    p.drawString(50, y, f"Comprobación: F(x) = {obj.comprobacion}")
    y -= 40

    p.drawString(50, y, "Iteraciones:")
    y -= 20
    for linea in obj.iteraciones.split('\n'):
        if y < 50:
            p.showPage()
            y = 800
        p.drawString(60, y, linea)
        y -= 20

    p.showPage()
    p.save()
    return response

@login_required
def repetir_punto_fijo(request, id):
    """
    Permite repetir un cálculo de Punto Fijo guardado en el historial.
    Solo muestra el proceso completo a usuarios premium.
    """
    obj = get_object_or_404(PuntoFijoHistorial, id=id)

    form = PuntoFijoForm(initial={
        'funcion': obj.funcion,
        'despeje': obj.despeje,
        'valor_inicial': obj.valor_inicial,
        'tolerancia': obj.tolerancia,
        'decimales': obj.decimales,
    })

    resultado = []
    solucion = None
    error = None
    comprobacion = None

    try:
        fx_input = obj.funcion
        gx_input = obj.despeje
        x0 = obj.valor_inicial
        tolerancia = obj.tolerancia
        decimales = obj.decimales

        i = 0
        ea = 100
        x1 = None

        while True:
            i += 1
            x1 = float(ne.evaluate(gx_input, local_dict={'x': x0}))
            fxi = float(ne.evaluate(fx_input, local_dict={'x': x1}))

            if i > 1:
                ea = abs((x1 - x0) / x1) * 100
            else:
                ea = None

            valor_x0 = round(x0, decimales)
            valor_x1 = round(x1, decimales)
            detalle = f"x{i} = {gx_input.replace('x', str(valor_x0))} = {valor_x1}"
            if ea is not None:
                error_detalle = f"ε{i} = |({valor_x1} - {valor_x0}) / {valor_x1}| × 100% = {round(ea, 2)}%"
            else:
                error_detalle = "ε1 = -"

            resultado.append({
                'iteracion': i,
                'detalle': detalle,
                'error_detalle': error_detalle,
                'x': valor_x1,
                'error': round(ea, 2) if ea is not None else '-',
            })

            if ea is not None and ea <= tolerancia:
                break

            x0 = x1

        if resultado and x1 is not None:
            solucion = round(x1, decimales)
            error = round(ea, 5)
            try:
                comprobacion = round(float(ne.evaluate(fx_input, local_dict={'x': solucion})), 8)
            except:
                comprobacion = "Error al comprobar"

    except Exception as e:
        form.add_error(None, f"Error al evaluar la función o el despeje: {e}")

    # --- Datos para gráfica Punto Fijo en modo repetir ---
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
    if resultado and is_premium:
        grafica_punto_fijo = {
            'x': [r['iteracion'] for r in resultado],
            'y': [r['x'] for r in resultado],
            'error': [r['error'] for r in resultado],
        }
    else:
        grafica_punto_fijo = None
    # --- end datos gráfica ---

    return render(request, 'punto_fijo/puntoFijo.html', {
        'form': form,
        # Mostrar solo el resultado mínimo para no premium, igual que en trazador cubico
        'resultado': resultado if is_premium else None,
        'solucion': solucion,
        'error': error,
        'comprobacion': comprobacion,
        'fx': obj.funcion,
        'gx': obj.despeje,
        'fx_latex': obj.funcion_latex or obj.funcion,
        'gx_latex': obj.despeje_latex or obj.despeje,
        'is_premium': is_premium,
        'valor_inicial_val': obj.valor_inicial,
        'tolerancia_val': obj.tolerancia,
        'decimales_val': obj.decimales,
        'modo_repetir': True,
        'mostrar_resultado_minimo': True,  # Forzar mostrar el bloque de resultado mínimo
        'grafica_punto_fijo': grafica_punto_fijo,
    })


@login_required
def trazador_pdf(request, id):
    """
    Genera un PDF con el historial de un cálculo de Trazador Cúbico.
    """
    obj = SplineHistory.objects.get(id=id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="trazador_{obj.id}.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)
    y = 800

    p.drawString(50, y, f"TRAZADOR CÚBICO - Historial ID {obj.id}")
    y -= 30
    p.drawString(50, y, f"Puntos: {obj.puntos}")
    y -= 20
    p.drawString(50, y, f"x evaluado: {obj.x_valor}")
    y -= 20
    p.drawString(50, y, f"Resultado: {obj.resultado}")
    y -= 40

    p.drawString(50, y, "Polinomio usado:")
    y -= 20
    for linea in obj.polinomio_usado.split('\n'):
        if y < 50:
            p.showPage()
            y = 800
        p.drawString(60, y, linea)
        y -= 20

    y -= 20
    p.drawString(50, y, "Razonamiento:")
    y -= 20
    for linea in obj.razonamiento.split('\n'):
        if y < 50:
            p.showPage()
            y = 800
        p.drawString(60, y, linea)
        y -= 20

    p.showPage()
    p.save()
    return response

@login_required
def repetir_trazador(request, id):
    """
    Permite repetir un cálculo de Trazador Cúbico guardado en el historial.
    Solo muestra el proceso completo a usuarios premium.
    """
    obj = get_object_or_404(SplineHistory, id=id)

    form = SplineInputForm(initial={
        'points': obj.puntos,
        'x_value': obj.x_valor
    })

    result = None
    grafica_trazador = None
    try:
        points = parse_points(obj.puntos)
        result = natural_cubic_spline(points, obj.x_valor)
        # --- Datos para gráfica Trazador Cúbico en modo repetir ---
        is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
        if result and 'tabla_hm' in result and is_premium:
            x_vals = [p[0] for p in points]
            y_vals = [p[1] for p in points]
            polinomios = result.get('polinomios', [])
            grafica_trazador = {
                'x': x_vals,
                'y': y_vals,
                'polinomios': polinomios,
                'x_query': result.get('x_query'),
                'y_query': result.get('resultado'),
            }
        # Si no es premium, no se muestra gráfica
        # --- end datos gráfica ---
    except Exception as e:
        form.add_error(None, f"Error en los datos: {e}")
        grafica_trazador = None

    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
    return render(request, 'trazador_cubico/trazadorCubico.html', {
        'form': form,
        'result': result if is_premium else {
            'razonamiento': result['razonamiento'] if result else '',
            'polinomio_usado': result['polinomio_usado'] if result else '',
            'x_query': result['x_query'] if result and 'x_query' in result and result['x_query'] not in [None, ''] else None,
            'resultado': result['resultado'] if result and 'resultado' in result and result['resultado'] not in [None, ''] else None,
        },
        'is_premium': is_premium,
        'modo_repetir': True,
        'grafica_trazador': grafica_trazador,
    })


from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth import login, logout


# Create your views here.


def login_view(request):
    """Renderiza la página de login simple."""
    return render(request, 'paginas/login.html')

def inicio_sesion(request):
    """Procesa el inicio de sesión de usuario."""
    if request.method == 'POST':
        form = AuthenticationForm(request, request.POST)
        if form.is_valid():
            login(request, form.get_user())
            return redirect( 'index')
        else:
            return render(request, "paginas/inicio_sesion.html", {'form': form, 'error': 'credenciales incorrectas'})
    else:
        form = AuthenticationForm()
    return render(request, "paginas/inicio_sesion.html")

def registro(request):
    """Registro de usuario básico (no recomendado, usar registro_usuario)."""
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect( 'inicio_sesion')
        else:
            return render(request, "paginas/registro.html", {'form': form, 'error': 'credenciales incorrectas'})
    else:
        form = UserCreationForm()

    return render(request, "paginas/registro.html")

def cerrar_sesion(request):
    """Cierra la sesión del usuario actual."""
    logout(request)
    return redirect('login') 
@login_required
def index(request):
    """Página principal de la aplicación."""
    return render(request,'paginas/index.html') 
@login_required
def tienda(request):
    """Página de compra de premium."""
    is_premium = False
    if request.user.is_authenticated:
        try:
            is_premium = request.user.userprofile.is_premium
        except Exception:
            is_premium = False
    return render(request,'paginas/tienda.html', {'is_premium': is_premium})
@login_required
def documentacion_trazadores(request):
    """Documentación del método de Trazador Cúbico."""
    return render(request, 'trazador_cubico/documentacion_trazador.html')
@login_required
def documentacion_punto(request):
    """Documentación del método de Punto Fijo."""
    return render(request, 'punto_fijo/documentacion_puntoFijo.html')

@login_required
def comprar_premium(request):
    """Convierte al usuario en premium."""
    profile = UserProfile.objects.get(user=request.user)
    if not profile.is_premium:
        profile.is_premium = True
        profile.save()
    return redirect('tienda')

def registro_usuario(request):
    """Registro de usuario extendido con perfil personalizado."""
    if request.method == 'POST':
        form = RegistroUsuarioForm(request.POST, request.FILES)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('index')
    else:
        form = RegistroUsuarioForm()
    return render(request, 'paginas/registro_usuario.html', {'form': form})

@login_required
def perfil_usuario(request):
    """Muestra el perfil del usuario. Redirige a edición si faltan datos."""
    profile = request.user.userprofile
    # Si algún campo obligatorio está vacío, redirigir a editar perfil
    if not profile.nombre_completo or not profile.carrera or not profile.carnet or not profile.ciclo:
        return redirect('editar_perfil')
    return render(request, 'paginas/perfil_usuario.html', {'profile': profile})

@login_required
def editar_perfil(request):
    """Permite editar los datos del perfil de usuario."""
    profile = request.user.userprofile
    if request.method == 'POST':
        form = EditarPerfilForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            form.save()
            return redirect('perfil_usuario')
    else:
        form = EditarPerfilForm(instance=profile)
    return render(request, 'paginas/editar_perfil.html', {'form': form})
@login_required
def creditos(request):
    """Muestra los créditos del proyecto y el equipo de desarrollo."""
    integrantes = [
        {
            'nombre': 'Nombre 1',
            'rol': 'Diseño de interfaz, backend, métodos',
        },
        {
            'nombre': 'Nombre 2',
            'rol': 'Base de datos, seguridad, pruebas',
        },
        # Agrega más integrantes y roles aquí
    ]
    return render(request, 'paginas/creditos.html', {'integrantes': integrantes})
# ==================== MÉTODO SIMPLEX ====================

# Variable global para definir decimales en el simplex
decimal_places = 2

# Variable global para almacenar nombres de variables personalizadas
variables_nombres = ['x1', 'x2']

def format_value(value: float) -> str:
    """Formatea un valor flotante con el número de decimales especificado"""
    return f"{value:.{decimal_places}f}"

def parse_latex_to_numbers(funcion_latex, restricciones_latex):
    """
    Convierte las expresiones LaTeX de MathLive a listas numéricas
    """
    try:
        # Parsear función objetivo
        # Ejemplo: "20x_1 + 40x_2" -> [20, 40]
        objetivo = []
        
        # Limpiar la expresión y extraer coeficientes
        funcion_clean = funcion_latex.replace('\\cdot', '*').replace(' ', '')
        
        # Buscar patrones como 20x_1, -15x_2, etc.
        import re
        patron_coef = r'([+-]?\d*)\*?x_\{?(\d+)\}?'
        matches = re.findall(patron_coef, funcion_clean)
        
        # Crear diccionario con coeficientes
        coef_dict = {}
        for coef_str, var_num in matches:
            if coef_str in ['', '+']:
                coef = 1
            elif coef_str == '-':
                coef = -1
            else:
                coef = int(coef_str)
            coef_dict[int(var_num)] = coef
        
        # Crear lista ordenada de coeficientes
        max_var = max(coef_dict.keys()) if coef_dict else 1
        objetivo = [coef_dict.get(i+1, 0) for i in range(max_var)]
        
        # Parsear restricciones
        # Ejemplo: ["2x_1 + 3x_2 \\leq 110", "4x_1 + x_2 \\leq 130"] para maximizar
        # Ejemplo: ["2x_1 + 3x_2 \\geq 110", "4x_1 + x_2 \\geq 130"] para minimizar
        restricciones = []
        
        for restriccion in restricciones_latex:
            # Separar lado izquierdo del derecho
            operador_encontrado = None
            if '\\leq' in restriccion:
                lado_izq, lado_der = restriccion.split('\\leq')
                operador_encontrado = '<='
            elif '\\le' in restriccion:
                lado_izq, lado_der = restriccion.split('\\le')
                operador_encontrado = '<='
            elif '\\geq' in restriccion:
                lado_izq, lado_der = restriccion.split('\\geq')
                operador_encontrado = '>='
            elif '\\ge' in restriccion:
                lado_izq, lado_der = restriccion.split('\\ge')
                operador_encontrado = '>='
            elif '≤' in restriccion:
                lado_izq, lado_der = restriccion.split('≤')
                operador_encontrado = '<='
            elif '≥' in restriccion:
                lado_izq, lado_der = restriccion.split('≥')
                operador_encontrado = '>='
            else:
                continue
            
            # Limpiar y parsear lado izquierdo
            lado_izq = lado_izq.replace('\\cdot', '*').replace(' ', '')
            lado_der = lado_der.strip()
            
            # Extraer coeficientes del lado izquierdo
            matches_rest = re.findall(patron_coef, lado_izq)
            coef_rest_dict = {}
            
            for coef_str, var_num in matches_rest:
                if coef_str in ['', '+']:
                    coef = 1
                elif coef_str == '-':
                    coef = -1
                else:
                    coef = int(coef_str)
                coef_rest_dict[int(var_num)] = coef
            
            # Crear fila de restricción [coef_x1, coef_x2, ..., b]
            fila_restriccion = [coef_rest_dict.get(i+1, 0) for i in range(max_var)]
            
            # Agregar término independiente (lado derecho)
            try:
                rhs = float(lado_der)
                
                # Para restricciones >=, multiplicar por -1 para convertir a <=
                if operador_encontrado == '>=':
                    fila_restriccion = [-coef for coef in fila_restriccion]
                    rhs = -rhs
                
                fila_restriccion.append(rhs)
                restricciones.append(fila_restriccion)
            except ValueError:
                continue
        
        return objetivo, restricciones
        
    except Exception as e:
        print(f"Error en parsing: {e}")
        # Valores por defecto en caso de error
        return [20, 40], [[2, 3, 110], [4, 1, 130]]

def prepare_initial_table(obj: List[float], constraints: List[List[float]], variables_nombres: List[str] = None) -> Tuple[List[List[float]], List[str]]:
    """Prepara la tabla inicial del método simplex"""
    num_vars = len(obj)
    num_constraints = len(constraints)

    # Validar que no haya RHS negativos (indicaría necesidad de métodos avanzados)
    for i, row in enumerate(constraints):
        rhs = row[-1]
        if rhs < 0:
            raise ValueError(f"⚠️ MÉTODO SIMPLEX ESTÁNDAR NO APLICABLE\n\n"
                           f"La restricción {i+1} tiene un valor del lado derecho negativo ({rhs}).\n"
                           f"Esto indica que el problema requiere:\n\n"
                           f"💡 MÉTODOS AVANZADOS:\n"
                           f"• Método de las Dos Fases\n"
                           f"• Método de la Gran M\n"
                           f"• Variables artificiales\n\n"
                           f"📚 El método simplex estándar solo puede resolver problemas con:\n"
                           f"• Restricciones del tipo ≤ para maximización\n"
                           f"• Valores del lado derecho no negativos")

    matrix = []
    for i, row in enumerate(constraints):
        vars_part = row[:-1]
        rhs = row[-1]
        slack = [0] * num_constraints
        slack[i] = 1
        matrix.append(vars_part + slack + [rhs])

    z_row = [-c for c in obj] + [0] * (num_constraints + 1)
    matrix.append(z_row)

    # Usar nombres personalizados si están disponibles, sino usar nombres por defecto
    if variables_nombres and len(variables_nombres) >= num_vars:
        var_names = variables_nombres[:num_vars]
    else:
        var_names = [f"x{i+1}" for i in range(num_vars)]
    
    return matrix, var_names

def generate_simplex_solution(matrix: List[List[float]], var_names: List[str], obj: List[float], constraints: List[List[float]], tipo_objetivo: str = "Maximizar") -> dict:
    """
    Genera la solución completa del método simplex adaptada para Django
    """
    num_constraints = len(matrix) - 1
    num_vars = len(var_names)

    slack_vars = [f"s{i+1}" for i in range(num_constraints)]
    all_vars = var_names + slack_vars + ["B"] 
    table = pd.DataFrame(matrix, columns=all_vars)

    basis = slack_vars.copy()
    
    # Datos para el template
    resultado = {
        'obj_original': obj,
        'constraints_original': constraints,
        'var_names': var_names,
        'slack_vars': slack_vars,
        'iteraciones': [],
        'solucion_optima': {},
        'valor_z': 0,
        'tabla_inicial': None,
        'modelo_matematico': None,
        'modelo_holgura': None
    }
    
    # Modelo matemático original
    # Para minimización, necesitamos mostrar el objetivo original (sin negar)
    if tipo_objetivo == "Minimizar":
        # Para minimización, el objetivo original es el negativo de lo que usamos internamente
        obj_original = [-c for c in obj]  # obj ya está negado para minimización
        obj_expr = " + ".join(f"{format_value(c)}x_{{{i+1}}}" for i, c in enumerate(obj_original))
        objetivo_texto = f"\\text{{Minimizar }} Z = {obj_expr}"
        funcion_z_estandar = f"Z - ({obj_expr}) = 0"
    else:
        obj_expr = " + ".join(f"{format_value(c)}x_{{{i+1}}}" for i, c in enumerate(obj))
        objetivo_texto = f"\\text{{Maximizar }} Z = {obj_expr}"
        funcion_z_estandar = f"Z - ({obj_expr}) = 0"
    
    constraints_expr = []
    for i, row in enumerate(constraints):
        lhs = " + ".join(f"{format_value(val)}x_{{{j+1}}}" for j, val in enumerate(row[:-1]) if val != 0)
        rhs = format_value(row[-1])
        # Para minimización, mostrar las restricciones originales (>=)
        if tipo_objetivo == "Minimizar":
            # Las restricciones ya están convertidas a <= internamente, pero mostramos las originales >=
            # Necesitamos deshacer la conversión para mostrar
            lhs_original = " + ".join(f"{format_value(-val)}x_{{{j+1}}}" for j, val in enumerate(row[:-1]) if val != 0)
            rhs_original = format_value(-row[-1])
            constraints_expr.append(f"{lhs_original} \\geq {rhs_original}")
        else:
            constraints_expr.append(f"{lhs} \\leq {rhs}")
    
    resultado['modelo_matematico'] = {
        'funcion_objetivo': objetivo_texto,
        'funcion_z_estandar': funcion_z_estandar,
        'restricciones': constraints_expr
    }
    
    # Modelo con variables de holgura
    transformed_expr = []
    for i, row in enumerate(constraints):
        lhs = " + ".join(f"{format_value(val)}x_{{{j+1}}}" for j, val in enumerate(row[:-1]) if val != 0)
        lhs += f" + s_{{{i+1}}}"
        rhs = format_value(row[-1])
        transformed_expr.append(f"{lhs} = {rhs}")
    
    resultado['modelo_holgura'] = transformed_expr
    
    # Tabla inicial
    tabla_inicial_html = generar_tabla_html(table, basis, 0, var_names, var_names)
    resultado['tabla_inicial'] = tabla_inicial_html

    # Iteraciones del algoritmo simplex
    iteration = 0
    while True:
        # Crear datos de la iteración actual
        iteracion_data = {
            'numero': iteration,
            'tabla_html': generar_tabla_html(table, basis, iteration, var_names, var_names),
            'es_optima': False,
            'variable_entra': None,
            'variable_sale': None,
            'elemento_pivote': None,
            'operaciones': []
        }
        
        # Verificar si es solución óptima
        last_row = table.iloc[-1, :-1]
        pivot_col_name = last_row.idxmin()
        
        if table.at[len(table) - 1, pivot_col_name] >= 0:
            iteracion_data['es_optima'] = True
            resultado['iteraciones'].append(iteracion_data)
            break

        # Encontrar variable que entra y sale
        ratios = []
        for i in range(len(table) - 1):
            col_val = table.at[i, pivot_col_name]
            if col_val > 0:
                ratios.append(table.at[i, "B"] / col_val) 
            else:
                ratios.append(float('inf'))

        pivot_row = ratios.index(min(ratios))
        pivot_element = table.at[pivot_row, pivot_col_name]

        entering = pivot_col_name
        leaving = basis[pivot_row]
        
        iteracion_data['variable_entra'] = entering
        iteracion_data['variable_sale'] = leaving
        iteracion_data['elemento_pivote'] = format_value(pivot_element)
        
        # Normalización de fila pivote
        operacion_pivote = []
        new_pivot_row = []
        for col in table.columns:
            original_val = table.at[pivot_row, col]
            new_val = original_val / pivot_element
            new_pivot_row.append(new_val)
            operacion_pivote.append({
                'variable': col,
                'operacion': f"{format_value(original_val)} ÷ {format_value(pivot_element)} = {format_value(new_val)}"
            })
        table.iloc[pivot_row] = new_pivot_row
        
        iteracion_data['operaciones'].append({
            'tipo': 'normalizacion',
            'fila': f"F{pivot_row+1}",
            'detalles': operacion_pivote
        })
        
        # Actualización de otras filas
        for i in range(len(table)):
            if i != pivot_row:
                factor = table.at[i, pivot_col_name]
                if factor != 0:
                    operacion_fila = []
                    original_row = table.iloc[i].copy()
                    new_row = []
                    for col in table.columns:
                        old_val = original_row[col]
                        pivot_val = table.at[pivot_row, col]
                        result = old_val - factor * pivot_val
                        new_row.append(result)
                        operacion_fila.append({
                            'variable': col,
                            'operacion': f"{format_value(old_val)} - ({format_value(factor)} × {format_value(pivot_val)}) = {format_value(result)}"
                        })
                    table.iloc[i] = new_row
                    
                    iteracion_data['operaciones'].append({
                        'tipo': 'actualizacion',
                        'fila': f"F{i+1}",
                        'factor': format_value(factor),
                        'detalles': operacion_fila
                    })

        basis[pivot_row] = entering
        resultado['iteraciones'].append(iteracion_data)
        iteration += 1

    # Solución óptima con nombres personalizados
    solution = {}
    
    # Agregar variables de decisión (personalizadas)
    for i, var_personalizada in enumerate(variables_nombres):
        # Usar directamente el nombre personalizado que está en var_names
        if i < len(var_names) and var_names[i] in basis:
            idx = basis.index(var_names[i])
            solution[var_personalizada] = table.at[idx, "B"]
        else:
            solution[var_personalizada] = 0.0
    
    # Agregar variables de holgura
    for slack_var in slack_vars:
        if slack_var in basis:
            idx = basis.index(slack_var)
            solution[slack_var] = table.at[idx, "B"]
        else:
            solution[slack_var] = 0.0

    resultado['solucion_optima'] = solution
    
    # Para minimización, el valor de Z debe ser negado porque internamente trabajamos con -Z
    z_value = table.at[len(table) - 1, "B"]
    if tipo_objetivo == "Minimizar":
        z_value = -z_value  # Regresar al valor original
    
    resultado['valor_z'] = z_value
    resultado['tipo_objetivo'] = tipo_objetivo
    
    return resultado

def parse_latex_to_numbers_with_custom_vars(funcion_latex, restricciones_latex, variables_nombres):
    """
    Convierte las expresiones LaTeX de MathLive a listas numéricas usando nombres de variables personalizados
    """
    try:
        # Parsear función objetivo
        objetivo = []
        
        # Limpiar la expresión y extraer coeficientes
        funcion_clean = funcion_latex.replace('\\cdot', '*').replace(' ', '')
        
        # Crear un mapeo de variables personalizadas a índices
        var_mapping = {var: i+1 for i, var in enumerate(variables_nombres)}
        
        # Buscar patrones de variables personalizadas en la función objetivo
        coef_dict = {}
        
        # Primero, intentar buscar variables personalizadas directamente
        for var_name in variables_nombres:
            # Buscar patrones como 20var_name, -15var_name, etc.
            patron = rf'([+-]?\d*)\*?{re.escape(var_name)}(?!\w)'
            matches = re.findall(patron, funcion_clean)
            
            coef_sum = 0
            for coef_str in matches:
                if coef_str in ['', '+']:
                    coef_sum += 1
                elif coef_str == '-':
                    coef_sum += -1
                else:
                    coef_sum += int(coef_str)
            
            if coef_sum != 0:
                coef_dict[var_mapping[var_name]] = coef_sum
        
        # Si no se encontraron variables personalizadas, usar el patrón original x_1, x_2
        if not coef_dict:
            patron_coef = r'([+-]?\d*)\*?x_\{?(\d+)\}?'
            matches = re.findall(patron_coef, funcion_clean)
            
            for coef_str, var_num in matches:
                if coef_str in ['', '+']:
                    coef = 1
                elif coef_str == '-':
                    coef = -1
                else:
                    coef = int(coef_str)
                coef_dict[int(var_num)] = coef
        
        # Crear lista ordenada de coeficientes
        max_var = max(coef_dict.keys()) if coef_dict else len(variables_nombres)
        objetivo = [coef_dict.get(i+1, 0) for i in range(max_var)]
        
        # Parsear restricciones
        restricciones = []
        
        for restriccion in restricciones_latex:
            # Separar lado izquierdo del derecho
            operador_encontrado = None
            if '\\leq' in restriccion:
                lado_izq, lado_der = restriccion.split('\\leq')
                operador_encontrado = '<='
            elif '\\le' in restriccion:
                lado_izq, lado_der = restriccion.split('\\le')
                operador_encontrado = '<='
            elif '\\geq' in restriccion:
                lado_izq, lado_der = restriccion.split('\\geq')
                operador_encontrado = '>='
            elif '\\ge' in restriccion:
                lado_izq, lado_der = restriccion.split('\\ge')
                operador_encontrado = '>='
            elif '≤' in restriccion:
                lado_izq, lado_der = restriccion.split('≤')
                operador_encontrado = '<='
            elif '≥' in restriccion:
                lado_izq, lado_der = restriccion.split('≥')
                operador_encontrado = '>='
            else:
                continue
            
            # Limpiar y parsear lado izquierdo
            lado_izq = lado_izq.replace('\\cdot', '*').replace(' ', '')
            lado_der = lado_der.strip()
            
            # Extraer coeficientes del lado izquierdo usando variables personalizadas
            coef_rest_dict = {}
            
            # Buscar variables personalizadas
            for var_name in variables_nombres:
                patron = rf'([+-]?\d*)\*?{re.escape(var_name)}(?!\w)'
                matches_rest = re.findall(patron, lado_izq)
                
                coef_sum = 0
                for coef_str in matches_rest:
                    if coef_str in ['', '+']:
                        coef_sum += 1
                    elif coef_str == '-':
                        coef_sum += -1
                    else:
                        coef_sum += int(coef_str)
                
                if coef_sum != 0:
                    coef_rest_dict[var_mapping[var_name]] = coef_sum
            
            # Si no se encontraron variables personalizadas, usar patrón x_1, x_2
            if not coef_rest_dict:
                patron_coef = r'([+-]?\d*)\*?x_\{?(\d+)\}?'
                matches_rest = re.findall(patron_coef, lado_izq)
                
                for coef_str, var_num in matches_rest:
                    if coef_str in ['', '+']:
                        coef = 1
                    elif coef_str == '-':
                        coef = -1
                    else:
                        coef = int(coef_str)
                    coef_rest_dict[int(var_num)] = coef
            
            # Crear fila de restricción [coef_x1, coef_x2, ..., b]
            fila_restriccion = [coef_rest_dict.get(i+1, 0) for i in range(max_var)]
            
            # Agregar término independiente (lado derecho)
            try:
                rhs = float(lado_der)
                
                # Para restricciones >=, multiplicar por -1 para convertir a <=
                if operador_encontrado == '>=':
                    fila_restriccion = [-coef for coef in fila_restriccion]
                    rhs = -rhs
                
                fila_restriccion.append(rhs)
                restricciones.append(fila_restriccion)
            except ValueError:
                continue
        
        return objetivo, restricciones
        
    except Exception as e:
        print(f"Error en parsing con variables personalizadas: {e}")
        # Valores por defecto en caso de error
        return [20, 40], [[2, 3, 110], [4, 1, 130]]

def generate_simplex_solution_with_custom_vars(matrix: List[List[float]], var_names: List[str], obj: List[float], constraints: List[List[float]], tipo_objetivo: str = "Maximizar", variables_nombres: List[str] = None) -> dict:
    """
    Genera la solución completa del método simplex adaptada para Django con variables personalizadas
    """
    if variables_nombres is None:
        variables_nombres = [f"x{i+1}" for i in range(len(obj))]
    
    num_constraints = len(matrix) - 1
    num_vars = len(var_names)

    slack_vars = [f"s{i+1}" for i in range(num_constraints)]
    all_vars = var_names + slack_vars + ["B"] 
    table = pd.DataFrame(matrix, columns=all_vars)

    basis = slack_vars.copy()
    
    # Datos para el template
    resultado = {
        'obj_original': obj,
        'constraints_original': constraints,
        'var_names': var_names,
        'slack_vars': slack_vars,
        'iteraciones': [],
        'solucion_optima': {},
        'valor_z': 0,
        'tabla_inicial': None,
        'modelo_matematico': None,
        'modelo_holgura': None,
        'variables_nombres': variables_nombres
    }
    
    # Modelo matemático original con variables personalizadas
    if tipo_objetivo == "Minimizar":
        # Para minimización, el objetivo original es el negativo de lo que usamos internamente
        obj_original = [-c for c in obj]  # obj ya está negado para minimización
        obj_expr = " + ".join(f"{format_value(c)}{var}" for i, (c, var) in enumerate(zip(obj_original, variables_nombres)) if c != 0)
        objetivo_texto = f"\\text{{Minimizar }} Z = {obj_expr}"
        funcion_z_estandar = f"Z - ({obj_expr}) = 0"
    else:
        obj_expr = " + ".join(f"{format_value(c)}{var}" for i, (c, var) in enumerate(zip(obj, variables_nombres)) if c != 0)
        objetivo_texto = f"\\text{{Maximizar }} Z = {obj_expr}"
        funcion_z_estandar = f"Z - ({obj_expr}) = 0"
    
    constraints_expr = []
    for i, row in enumerate(constraints):
        lhs = " + ".join(f"{format_value(val)}{variables_nombres[j]}" for j, val in enumerate(row[:-1]) if val != 0)
        rhs = format_value(row[-1])
        # Para minimización, mostrar las restricciones originales (>=)
        if tipo_objetivo == "Minimizar":
            # Las restricciones ya están convertidas a <= internamente, pero mostramos las originales >=
            lhs_original = " + ".join(f"{format_value(-val)}{variables_nombres[j]}" for j, val in enumerate(row[:-1]) if val != 0)
            rhs_original = format_value(-row[-1])
            constraints_expr.append(f"{lhs_original} \\geq {rhs_original}")
        else:
            constraints_expr.append(f"{lhs} \\leq {rhs}")
    
    resultado['modelo_matematico'] = {
        'funcion_objetivo': objetivo_texto,
        'funcion_z_estandar': funcion_z_estandar,
        'restricciones': constraints_expr
    }
    
    # Modelo con variables de holgura
    transformed_expr = []
    for i, row in enumerate(constraints):
        lhs = " + ".join(f"{format_value(val)}{variables_nombres[j]}" for j, val in enumerate(row[:-1]) if val != 0)
        lhs += f" + s_{{{i+1}}}"
        rhs = format_value(row[-1])
        transformed_expr.append(f"{lhs} = {rhs}")
    
    resultado['modelo_holgura'] = transformed_expr
    
    # Tabla inicial
    tabla_inicial_html = generar_tabla_html(table, basis, 0, var_names, var_names)
    resultado['tabla_inicial'] = tabla_inicial_html

    # Iteraciones del algoritmo simplex
    iteration = 0
    while True:
        # Crear datos de la iteración actual
        iteracion_data = {
            'numero': iteration,
            'tabla_html': generar_tabla_html(table, basis, iteration, var_names, var_names),
            'es_optima': False,
            'variable_entra': None,
            'variable_sale': None,
            'elemento_pivote': None,
            'operaciones': []
        }
        
        # Verificar si es solución óptima
        last_row = table.iloc[-1, :-1]
        pivot_col_name = last_row.idxmin()
        
        if table.at[len(table) - 1, pivot_col_name] >= 0:
            iteracion_data['es_optima'] = True
            resultado['iteraciones'].append(iteracion_data)
            break

        # Encontrar variable que entra y sale
        ratios = []
        for i in range(len(table) - 1):
            col_val = table.at[i, pivot_col_name]
            if col_val > 0:
                ratios.append(table.at[i, "B"] / col_val) 
            else:
                ratios.append(float('inf'))

        pivot_row = ratios.index(min(ratios))
        pivot_element = table.at[pivot_row, pivot_col_name]

        entering = pivot_col_name
        leaving = basis[pivot_row]
        
        iteracion_data['variable_entra'] = entering
        iteracion_data['variable_sale'] = leaving
        iteracion_data['elemento_pivote'] = format_value(pivot_element)
        
        # Normalización de fila pivote
        operacion_pivote = []
        new_pivot_row = []
        for col in table.columns:
            original_val = table.at[pivot_row, col]
            new_val = original_val / pivot_element
            new_pivot_row.append(new_val)
            operacion_pivote.append({
                'variable': col,
                'operacion': f"{format_value(original_val)} ÷ {format_value(pivot_element)} = {format_value(new_val)}"
            })
        table.iloc[pivot_row] = new_pivot_row
        
        iteracion_data['operaciones'].append({
            'tipo': 'normalizacion',
            'fila': f"F{pivot_row+1}",
            'detalles': operacion_pivote
        })
        
        # Actualización de otras filas
        for i in range(len(table)):
            if i != pivot_row:
                factor = table.at[i, pivot_col_name]
                if factor != 0:
                    operacion_fila = []
                    original_row = table.iloc[i].copy()
                    new_row = []
                    for col in table.columns:
                        old_val = original_row[col]
                        pivot_val = table.at[pivot_row, col]
                        result = old_val - factor * pivot_val
                        new_row.append(result)
                        operacion_fila.append({
                            'variable': col,
                            'operacion': f"{format_value(old_val)} - ({format_value(factor)} × {format_value(pivot_val)}) = {format_value(result)}"
                        })
                    table.iloc[i] = new_row
                    
                    iteracion_data['operaciones'].append({
                        'tipo': 'actualizacion',
                        'fila': f"F{i+1}",
                        'factor': format_value(factor),
                        'detalles': operacion_fila
                    })

        basis[pivot_row] = entering
        resultado['iteraciones'].append(iteracion_data)
        iteration += 1

    # Solución óptima con nombres personalizados
    solution = {}
    
    # Agregar variables de decisión (personalizadas)
    for i, var_personalizada in enumerate(variables_nombres):
        # Usar directamente el nombre personalizado que está en var_names
        if i < len(var_names) and var_names[i] in basis:
            idx = basis.index(var_names[i])
            solution[var_personalizada] = table.at[idx, "B"]
        else:
            solution[var_personalizada] = 0.0
    
    # Agregar variables de holgura
    for slack_var in slack_vars:
        if slack_var in basis:
            idx = basis.index(slack_var)
            solution[slack_var] = table.at[idx, "B"]
        else:
            solution[slack_var] = 0.0

    resultado['solucion_optima'] = solution
    
    # Para minimización, el valor de Z debe ser negado porque internamente trabajamos con -Z
    z_value = table.at[len(table) - 1, "B"]
    if tipo_objetivo == "Minimizar":
        z_value = -z_value  # Regresar al valor original
    
    resultado['valor_z'] = z_value
    resultado['tipo_objetivo'] = tipo_objetivo
    
    return resultado

def convertir_nombre_variable_a_personalizado(nombre_variable, var_names, variables_nombres):
    """
    Convierte un nombre de variable del DataFrame a su nombre personalizado
    """
    if variables_nombres and nombre_variable in var_names:
        # Encontrar el índice en var_names y obtener el nombre personalizado
        try:
            index = var_names.index(nombre_variable)
            if index < len(variables_nombres):
                return variables_nombres[index]
        except ValueError:
            pass
    return nombre_variable

def generar_tabla_html(table, basis, iteration_num, variables_nombres=None, var_names=None):
    """Genera HTML para una tabla del simplex con nombres de variables personalizados"""
    html = f'<div class="tabla-iteracion">'
    html += f'<h4>Iteración {iteration_num}</h4>'
    html += '<table class="simplex-table">'
    html += '<thead><tr><th>Base</th>'
    
    # Usar nombres personalizados para las columnas si están disponibles
    columnas = list(table.columns)
    for i, col in enumerate(columnas):
        if variables_nombres and var_names and col in var_names:
            # Convertir a nombre personalizado
            display_name = convertir_nombre_variable_a_personalizado(col, var_names, variables_nombres)
        else:
            display_name = col
        html += f'<th>{display_name}</th>'
    html += '</tr></thead><tbody>'

    for i in range(len(table)):
        html += '<tr>'
        # Para la base también usar nombres personalizados si es una variable original
        base_name = basis[i] if i < len(basis) else 'Z'
        if variables_nombres and var_names and base_name in var_names:
            base_display = convertir_nombre_variable_a_personalizado(base_name, var_names, variables_nombres)
        else:
            base_display = base_name
        html += f'<td><strong>{base_display}</strong></td>'
        
        for col in table.columns:
            val = table.at[i, col]
            html += f'<td>{format_value(val)}</td>'
        html += '</tr>'
    html += '</tbody></table></div>'
    
    return html

def validar_coherencia_objetivo_restricciones(tipo_objetivo, restricciones):
    """
    Valida que el tipo de objetivo sea coherente con las restricciones:
    - Maximizar: solo permite restricciones <=
    - Minimizar: solo permite restricciones >=
    
    IMPORTANTE: El método simplex estándar implementado NO puede resolver:
    - Problemas de minimización (requieren Método de las Dos Fases o Gran M)
    - Problemas con restricciones mixtas
    - Problemas con RHS negativos
    """
    for restriccion in restricciones:
        if tipo_objetivo == "Maximizar":
            # Para maximizar, no se permiten restricciones >=
            if "\\geq" in restriccion or "\\ge" in restriccion:
                return ("⚠️ PROBLEMA NO VÁLIDO PARA MÉTODO SIMPLEX ESTÁNDAR\n\n"
                       "Para problemas de MAXIMIZACIÓN solo se permiten restricciones del tipo ≤ (menor o igual).\n"
                       "Su problema contiene restricciones ≥ (mayor o igual).\n\n"
                       "💡 SOLUCIÓN: Debe aplicar el Método de la Gran M o Método de las Dos Fases para resolver este problema.")
        else:  # Minimizar
            # IMPORTANTE: Los problemas de minimización requieren métodos avanzados
            return ("⚠️ PROBLEMA DE MINIMIZACIÓN DETECTADO\n\n"
                   "Los problemas de MINIMIZACIÓN con restricciones ≥ (mayor o igual) requieren métodos avanzados que no están implementados en esta versión del sistema.\n\n"
                   "💡 SOLUCIÓN RECOMENDADA:\n"
                   "• Use el Método de las Dos Fases\n"
                   "• Use el Método de la Gran M\n"
                   "• Convierta manualmente a un problema de maximización equivalente\n\n"
                   "� ALTERNATIVA MANUAL:\n"
                   "1. Cambie 'Minimizar Z' por 'Maximizar W = -Z'\n"
                   "2. Multiplique la función objetivo por -1\n"
                   "3. Agregue variables de holgura y artificiales según corresponda")
    
    return None  # No hay errores

@login_required
def simplex(request):
    """Vista principal del método simplex"""
    if request.method == "POST":
        try:
            # Obtener datos del formulario
            funcion_objetivo = request.POST.get("funcion_objetivo", "")
            restricciones_raw = request.POST.get("restricciones", "[]")
            tipo_objetivo = request.POST.get("tipo_objetivo", "Maximizar")
            variables_nombres_raw = request.POST.get("variables_nombres", "[]")
            
            # Parsear las restricciones JSON y variables personalizadas
            restricciones = json.loads(restricciones_raw)
            variables_nombres = json.loads(variables_nombres_raw) if variables_nombres_raw else []
            
            # Si no hay variables personalizadas, usar valores por defecto
            if not variables_nombres:
                variables_nombres = ['x1', 'x2']
            
            # Validar coherencia entre tipo de objetivo y restricciones (validación adicional del backend)
            error_validacion = validar_coherencia_objetivo_restricciones(tipo_objetivo, restricciones)
            if error_validacion:
                return render(request, 'simplex/simplex.html', {
                    'error': error_validacion
                })
            
            # Convertir de LaTeX a números usando variables personalizadas
            objetivo, constraints = parse_latex_to_numbers_with_custom_vars(funcion_objetivo, restricciones, variables_nombres)
            
            # Para minimización, convertir a maximización (multiplicar objetivo por -1)
            if tipo_objetivo == "Minimizar":
                objetivo = [-coef for coef in objetivo]
            
            # Preparar tabla inicial con variables personalizadas
            matrix, var_names = prepare_initial_table(objetivo, constraints, variables_nombres)
            
            # Resolver el problema simplex con variables personalizadas
            resultado = generate_simplex_solution_with_custom_vars(matrix, var_names, objetivo, constraints, tipo_objetivo, variables_nombres)
            
            # Verificar si el usuario es premium
            is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium if request.user.is_authenticated else False
            
            # Guardar en historial si el usuario está autenticado
            if request.user.is_authenticated and resultado:
                SimplexHistorial.objects.create(
                    user=request.user,
                    tipo_objetivo=tipo_objetivo,
                    funcion_objetivo=funcion_objetivo,
                    restricciones=restricciones_raw,
                    solucion_optima=json.dumps(resultado.get('solucion_optima', {})),
                    valor_z=resultado.get('valor_z', 0),
                    iteraciones_json=json.dumps(resultado.get('iteraciones', [])) if is_premium else "[]",
                    modelo_matematico=json.dumps(resultado.get('modelo_matematico', {}))
                )
            
            return render(request, 'simplex/simplex.html', {
                'funcion_objetivo': funcion_objetivo,
                'restricciones': restricciones,
                'resultado': resultado,
                'tiene_solucion': True,
                'tipo_objetivo': tipo_objetivo,
                'is_premium': is_premium,
                'variables_nombres': variables_nombres
            })
            
        except Exception as e:
            print(f"Error en simplex: {e}")
            return render(request, 'simplex/simplex.html', {
                'error': f"Error al procesar los datos: {str(e)}"
            })

    return render(request, 'simplex/simplex.html')

@login_required
def simplex_historial(request):
    """Vista para mostrar el historial del método simplex del usuario"""
    
    historial = SimplexHistorial.objects.filter(user=request.user).order_by('-fecha')
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium
    
    # Parsear JSON de solucion_optima para cada item del historial
    for item in historial:
        try:
            item.solucion_optima = json.loads(item.solucion_optima)
        except:
            item.solucion_optima = {}
    
    return render(request, 'simplex/historial.html', {        
        'historial': historial,
        'is_premium': is_premium
    })

from django.shortcuts import render
from sympy import symbols, Eq, solve, sympify, Float, Number
import numpy as np
import plotly.graph_objects as go
from itertools import combinations
from math import atan2
import traceback
from .models import MetodoGraficoHistorial
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter



def safe_float(value):
    """Convierte cualquier número a float nativo de Python"""
    if isinstance(value, (Float, Number)):
        return float(value.evalf())
    try:
        return float(value)
    except:
        return 0.0

def parse_expression(expr_str, variables):
    """Analiza una expresión matemática y devuelve sus componentes"""
    try:
        expr = sympify(expr_str)
        
        if expr.is_Number:
            return {
                'expression': expr,
                'coefficients': {var: 0.0 for var in variables},
                'constant': safe_float(expr),
                'original': expr_str
            }
        
        coeffs = {}
        for var in variables:
            coeff = expr.coeff(var)
            coeffs[var] = safe_float(coeff) if coeff is not None else 0.0
            
        constant = safe_float(expr.subs({var: 0 for var in variables}))
        
        return {
            'expression': expr,
            'coefficients': coeffs,
            'constant': constant,
            'original': expr_str
        }
    except Exception as e:
        raise ValueError(f"Error al analizar expresión: {expr_str}. {str(e)}")

def metodo_grafico(request):
    x, y = symbols('x y')
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium
    # Valores por defecto
    default_objective = '3*x + 2*y'
    default_restrictions = '2*x + y <= 100\nx + y <= 80'
    default_optimization = 'max'
    context = {
        'default_function': default_objective,
        'default_restrictions': default_restrictions,
        'default_optimization': default_optimization,
        'steps': [],
        'errors': [],
        'show_results': False,
        'is_premium': is_premium,
        'show_graph': False
    }

    # Inicializar variables para los inputs
    objective = default_objective
    optimization = default_optimization
    restrictions_text = default_restrictions

    if request.method == 'POST':
        # Tomar SIEMPRE los valores ingresados por el usuario
        objective = request.POST.get('objective', default_objective).strip()
        optimization = request.POST.get('optimization', default_optimization)
        restrictions_text = request.POST.get('restrictions', default_restrictions).strip()
        context.update({
            'objective': objective,
            'optimization': optimization,
            'restrictions_text': restrictions_text
        })

        try:
            # Paso 1: Obtener datos de entrada
            try:
                objective_expr = parse_expression(objective, [x, y])
                coeff_x = safe_float(objective_expr['coefficients'][x])
                coeff_y = safe_float(objective_expr['coefficients'][y])
                
                # Construir procedimiento matemático
                math_steps = []
                math_steps.append({
                    'title': '1. Función Objetivo',
                    'content': f"Z = {coeff_x}x + {coeff_y}y",
                    'optimization': f"Optimizar: {'Máximo' if optimization == 'max' else 'Mínimo'}"
                })
            except ValueError as e:
                context['errors'].append(str(e))
                return render(request, 'metodo_grafico/metodo_grafico.html', context)

            # Paso 3: Procesar restricciones
            restrictions = []
            restrictions_list = []
            for i, line in enumerate(restrictions_text.split('\n')):
                line = line.strip()
                if not line:
                    continue
                
                try:
                    if '<=' in line:
                        left, right = line.split('<=')
                        op = '<='
                    elif '>=' in line:
                        left, right = line.split('>=')
                        op = '>='
                    else:
                        raise ValueError("Operador no soportado. Use <= o >=")
                    
                    left_expr = parse_expression(left.strip(), [x, y])
                    right_expr = parse_expression(right.strip(), [x, y])
                    
                    restrictions.append({
                        'left': left_expr,
                        'right': right_expr,
                        'op': op,
                        'display': f"{left.strip()} {op} {right.strip()}"
                    })
                    restrictions_list.append(f"{left.strip()} {op} {right.strip()}")
                    
                except Exception as e:
                    context['errors'].append(f"Error en restricción {i+1}: {str(e)}")
                    return render(request, 'metodo_grafico/metodo_grafico.html', context)

            # Restricciones de no negatividad
            restrictions.extend([
                {
                    'left': parse_expression('x', [x, y]),
                    'right': parse_expression('0', [x, y]),
                    'op': '>=',
                    'display': 'x >= 0'
                },
                {
                    'left': parse_expression('y', [x, y]),
                    'right': parse_expression('0', [x, y]),
                    'op': '>=',
                    'display': 'y >= 0'
                }
            ])
            restrictions_list.extend(['x >= 0', 'y >= 0'])
            
            math_steps.append({
                'title': '2. Restricciones',
                'items': restrictions_list
            })

            # Paso 4: Calcular intersecciones con detalles
            equations = [Eq(r['left']['expression'], r['right']['expression']) for r in restrictions]
            vertices = []
            intersection_details = []
            
            for i, (eq1, eq2) in enumerate(combinations(equations, 2), 1):
                try:
                    sol = solve((eq1, eq2), (x, y), dict=True)
                    if sol and x in sol[0] and y in sol[0]:
                        x_val = safe_float(sol[0][x])
                        y_val = safe_float(sol[0][y])
                        
                        if x_val >= -1e-6 and y_val >= -1e-6:
                            x_val, y_val = max(x_val, 0), max(y_val, 0)
                            vertices.append((x_val, y_val))
                            
                            # Detalles de solución
                            a1 = safe_float(eq1.lhs.coeff(x))
                            b1 = safe_float(eq1.lhs.coeff(y))
                            c1 = safe_float(eq1.rhs)
                            
                            a2 = safe_float(eq2.lhs.coeff(x))
                            b2 = safe_float(eq2.lhs.coeff(y))
                            c2 = safe_float(eq2.rhs)
                            
                            solution_steps = [
                                f"Intersección entre:",
                                f"1) {a1}x + {b1}y = {c1}",
                                f"2) {a2}x + {b2}y = {c2}",
                                f"Resolviendo el sistema:"
                            ]
                            
                            # Método de reducción
                            solution_steps.append("Método de reducción:")
                            
                            # Multiplicar ecuaciones para igualar coeficientes
                            mult1 = a2
                            mult2 = a1
                            
                            eq1_mult = f"{mult1}*(1): {mult1*a1}x + {mult1*b1}y = {mult1*c1}"
                            eq2_mult = f"{mult2}*(2): {mult2*a2}x + {mult2*b2}y = {mult2*c2}"
                            
                            solution_steps.append(eq1_mult)
                            solution_steps.append(eq2_mult)
                            
                            # Restar ecuaciones
                            new_eq = f"({mult1*b1} - {mult2*b2})y = {mult1*c1} - {mult2*c2}"
                            solution_steps.append(f"Restando: {new_eq}")
                            
                            # Resolver para y
                            y_coeff = mult1*b1 - mult2*b2
                            y_rhs = mult1*c1 - mult2*c2
                            y_sol = f"y = {y_rhs}/{y_coeff} = {y_val:.4f}"
                            solution_steps.append(y_sol)
                            
                            # Sustituir y para encontrar x
                            x_sol = f"Sustituyendo y en (1): x = ({c1} - {b1}*{y_val:.4f})/{a1} = {x_val:.4f}"
                            solution_steps.append(x_sol)
                            
                            solution_steps.append(f"Punto de intersección: ({x_val:.4f}, {y_val:.4f})");
                            
                            intersection_details.append({
                                'equations': [f"{a1}x + {b1}y = {c1}", f"{a2}x + {b2}y = {c2}"],
                                'steps': solution_steps,
                                'point': (x_val, y_val)
                            })
                except Exception as e:
                    context['errors'].append(f"Error al calcular intersección: {str(e)}")
                    continue

            # Eliminar duplicados
            vertices = list(set([(round(x, 4), round(y, 4)) for x, y in vertices]))
            
            math_steps.append({
                'title': '3. Puntos de Intersección',
                'intersections': intersection_details
            })

            # Paso 5: Validación detallada de puntos factibles
            feasible_points = []
            feasibility_details = []
            
            for i, (vx, vy) in enumerate(vertices, 1):
                is_feasible = True
                validation_steps = []
                
                for r in restrictions:
                    left_val = safe_float(r['left']['expression'].subs({x: vx, y: vy}))
                    right_val = safe_float(r['right']['expression'].subs({x: vx, y: vy}))
                    op = r['op']
                    
                    if op == '<=':
                        satisfies = left_val <= right_val + 1e-6
                    elif op == '>=':
                        satisfies = left_val >= right_val - 1e-6
                    
                    validation_steps.append(
                        f"{r['display']}: {left_val:.2f} {op} {right_val:.2f} {'✔️' if satisfies else '❌'}"
                    )
                    
                    if not satisfies:
                        is_feasible = False
                
                if is_feasible:
                    feasible_points.append((vx, vy))
                    feasibility_details.append({
                        'id': len(feasible_points),
                        'x': vx,
                        'y': vy,
                        'validations': validation_steps
                    })

            if not feasible_points:
                context['errors'].append("No se encontraron puntos factibles que satisfagan todas las restricciones")
                return render(request, 'metodo_grafico/metodo_grafico.html', context)
            
            math_steps.append({
                'title': '4. Puntos Factibles',
                'points': feasibility_details
            })

            # Paso 6: Evaluar función objetivo
            solutions = []
            evaluation_steps = []
            
            for i, point in enumerate(feasibility_details, 1):


                vx, vy = point['x'], point['y']
                z = coeff_x * vx + coeff_y * vy
                solutions.append({
                    'id': i,
                    'x': vx,
                    'y': vy,
                    'z': z,
                    'is_optimal': False
                })
                
                evaluation_steps.append(
                    f"V{i}: Z = {coeff_x}×{vx:.4f} + {coeff_y}×{vy:.4f} = {z:.4f}"
                )
            
            math_steps.append({
                'title': '5. Evaluación de la Función Objetivo',
                'evaluations': evaluation_steps
            })

            # Paso 7: Encontrar solución óptima
            if optimization == 'max':
                optimal = max(solutions, key=lambda p: p['z'])
            else:
                optimal = min(solutions, key=lambda p: p['z'])
            
            for sol in solutions:
                sol['is_optimal'] = abs(sol['z'] - optimal['z']) < 1e-6
            
            context['optimal'] = optimal
            context['solutions'] = solutions
            context['show_results'] = True
            context['math_steps'] = math_steps
            
            math_steps.append({
                'title': '6. Solución Óptima',
                'point': f"V{optimal['id']} ({optimal['x']:.4f}, {optimal['y']:.4f})",
                'value': f"Z = {optimal['z']:.4f}",
                'type': optimization
            })

            # Preparar datos para el historial
            puntos_factibles_str = "\n".join([f"V{i}: ({p['x']:.4f}, {p['y']:.4f})" for i, p in enumerate(feasibility_details, 1)])
            vertices_str = "\n".join([f"({x:.4f}, {y:.4f})" for x, y in vertices])
            def paso_a_texto(step):
                partes = [f"{step.get('title','')}:"]
                if step.get('content'):
                    partes.append(str(step['content']))
                if step.get('items'):
                    partes.extend([str(item) for item in step['items']])
                if step.get('evaluations'):
                    partes.extend([str(ev) for ev in step['evaluations']])
                if step.get('intersections'):
                    for inter in step['intersections']:
                        if 'steps' in inter:
                            partes.extend([str(s) for s in inter['steps']])
                if step.get('points'):
                    for pt in step['points']:
                        if 'x' in pt and 'y' in pt:
                            partes.append(f"Punto: ({pt['x']:.4f}, {pt['y']:.4f})")
                        if 'validations' in pt:
                            partes.extend([str(val) for val in pt['validations']])
                if step.get('point'):
                    partes.append(str(step['point']))
                if step.get('value'):
                    partes.append(str(step['value']))
                if step.get('type'):
                    partes.append(str(step['type']))
                return '\n'.join(partes)
            iteraciones_str = '\n\n'.join([paso_a_texto(step) for step in math_steps])

            if is_premium and context['show_results']:
                try:
                    # Generar la gráfica con ajustes visuales
                    fig = go.Figure()
                    
                    # Configurar ejes y rangos
                    x_vals = [p['x'] for p in solutions] + [0]
                    y_vals = [p['y'] for p in solutions] + [0]
                    x_max = max(x_vals) + 10
                    y_max = max(y_vals) + 10
                    
                    # Dibujar restricciones con estilo mejorado
                    for r in restrictions[:-2]:  # Excluir x>=0, y>=0
                        try:
                            a = safe_float(r['left']['coefficients'].get(x, 0))
                            b = safe_float(r['left']['coefficients'].get(y, 0))
                            c = safe_float(r['right']['expression'])
                            
                            if b != 0:
                                x_points = np.linspace(0, x_max, 100)
                                y_points = [(c - a * x) / b for x in x_points]
                                fig.add_trace(go.Scatter(
                                    x=x_points,
                                    y=y_points,
                                    mode='lines',
                                    line=dict(color='blue', width=1.5, dash='dot'),
                                    name=r['display'],
                                    hoverinfo='text',
                                    hovertext=r['display']
                                ))
                        except:
                            continue
                    
                    # Dibujar región factible (convex hull)
                    if len(feasible_points) >= 3:
                        from scipy.spatial import ConvexHull
                        hull = ConvexHull(np.array(feasible_points))
                        sorted_points = [feasible_points[i] for i in hull.vertices]
                        
                        x_region = [p[0] for p in sorted_points] + [sorted_points[0][0]]
                        y_region = [p[1] for p in sorted_points] + [sorted_points[0][1]]
                        
                        fig.add_trace(go.Scatter(
                            x=x_region,
                            y=y_region,
                            fill='toself',
                            fillcolor='rgba(100, 200, 100, 0.3)',
                            line=dict(color='green', width=2),
                            name='Región Factible',
                            hoverinfo='none'
                        ))
                    
                    # Marcar puntos factibles
                    fig.add_trace(go.Scatter(
                        x=[p['x'] for p in solutions],
                        y=[p['y'] for p in solutions],
                        mode='markers+text',
                        text=[f"V{p['id']}" for p in solutions],
                        textposition="top center",
                        marker=dict(size=10, color='red'),
                        name='Vértices',
                        hoverinfo='text',
                        hovertext=[f"V{p['id']}: ({p['x']:.2f}, {p['y']:.2f})" for p in solutions]
                    ))
                    
                    # Marcar solución óptima
                    fig.add_trace(go.Scatter(
                        x=[optimal['x']],
                        y=[optimal['y']],
                        mode='markers+text',
                        text=['ÓPTIMO'],
                        textposition="top center",
                        marker=dict(size=16, color='gold', symbol='star'),
                        name='Óptimo',
                        hoverinfo='text',
                        hovertext=f"Óptimo: ({optimal['x']:.2f}, {optimal['y']:.2f})<br>Z = {optimal['z']:.2f}"
                    ))
                    
                    # Configuración del layout mejorado
                    fig.update_layout(
                        title='Solución del Método Gráfico',
                        xaxis_title='Variable x',
                        yaxis_title='Variable y',
                        xaxis=dict(
                            range=[0, x_max],
                            showgrid=True,
                            gridcolor='lightgray',
                            zeroline=True,
                            zerolinecolor='black'
                        ),
                        yaxis=dict(
                            range=[0, y_max],
                            showgrid=True,
                            gridcolor='lightgray',
                            zeroline=True,
                            zerolinecolor='black'
                        ),
                        plot_bgcolor='white',
                        height=600,
                        width=800,
                        margin=dict(l=50, r=50, b=100, t=100, pad=4),
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="right",
                            x=1
                        )
                    )
                    
                    # Convertir la figura a HTML
                    plot_html = fig.to_html(full_html=False, include_plotlyjs='cdn')
                    context['graph'] = plot_html
                    context['show_graph'] = True
                    # --- GUARDAR PNG TEMPORAL ---
                    import tempfile
                    timestamp = int(time.time())
                    user_id = request.user.id if request.user.is_authenticated else 'anon'
                    png_filename = f"grafica_{user_id}_{timestamp}.png"
                    png_path = os.path.join('calculadora', 'static', 'graficas_tmp', png_filename)
                    fig.write_image(png_path)
                    context['png_path'] = png_path  # Para el PDF
                except Exception as e:
                    print(f"Error al generar gráfica: {str(e)}")
                    context['errors'].append("Error al generar la gráfica")
                    plot_html = None
                    context['png_path'] = None
            # Guardar en el historial
            MetodoGraficoHistorial.objects.create(
                user=request.user,
                funcion=objective,
                optimizacion=optimization,
                restricciones=restrictions_text,
                solucion=f"Z = {optimal['z']:.4f}",
                punto_optimo=f"({optimal['x']:.4f}, {optimal['y']:.4f})",
                puntos_factibles=puntos_factibles_str,
                vertices=vertices_str,
                iteraciones=iteraciones_str,
                grafica=plot_html if is_premium and context['show_graph'] else None
                # No guardamos el PNG en el modelo, solo en contexto
            )
        except Exception as e:
            context['errors'].append(f"Error inesperado: {str(e)}")
            print(f"Error en metodo_grafico: {str(e)}")
    else:
        # GET: mostrar los valores por defecto SOLO si no hay datos previos
        context.update({
            'objective': objective,
            'optimization': optimization,
            'restrictions_text': restrictions_text
        })

    # Pasar procedimiento_texto al contexto
    if context.get('show_results'):
        context['procedimiento_texto'] = iteraciones_str
    else:
        context['procedimiento_texto'] = ''
    
    # Asegurar que historial_data siempre esté en el contexto para la plantilla
    if 'historial_data' not in context or not context['historial_data']:
        context['historial_data'] = {
            'objective': '',
            'optimization': '',
            'restrictions': '',
        }
    else:
        # Si falta alguna clave, la agregamos vacía
        for k in ['objective', 'optimization', 'restrictions']:
            if k not in context['historial_data']:
                context['historial_data'][k] = ''

    return render(request, 'metodo_grafico/metodo_grafico.html', context)

@login_required
def historial_metodo_grafico(request):
    """
    Muestra el historial de cálculos del Método Gráfico del usuario.
    """
    historial = MetodoGraficoHistorial.objects.filter(user=request.user).order_by('-fecha_creacion')
    is_premium = hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium
    
    return render(request, 'metodo_grafico/historialMetodoGrafico.html', {
        'historial': historial,
        'is_premium': is_premium
    })

@login_required
def documentacion_simplex(request):
    """Documentación del método de Simplex."""
    return render(request, 'simplex/documentacion_simplex.html')
@login_required
def documentacion_m(request):
    """Documentación del método la de M."""
    return render(request, 'simplex/documentacion_m.html')
@login_required
def documentacion_grafico(request):
    """Documentación del método de Grafico."""
    return render(request, 'simplex/documentacion_grafico.html')

@login_required
def cargar_simplex_historial(request, historial_id):
    """Cargar un problema del historial para resolverlo nuevamente"""
    
    try:
        historial_item = get_object_or_404(SimplexHistorial, id=historial_id, user=request.user)
        
        # Preparar datos para precargar el formulario
        # La función objetivo debe mantenerse tal como está guardada
        funcion_objetivo = historial_item.funcion_objetivo
        
        # Cargar restricciones tal como están guardadas
        restricciones_raw = json.loads(historial_item.restricciones)
        
        print(f"DEBUG - Restricciones cargadas del historial: {restricciones_raw}")
        
        contexto = {
            'precarga': {
                'tipo_objetivo': historial_item.tipo_objetivo,
                'funcion_objetivo': funcion_objetivo,
                'restricciones_json': json.dumps(restricciones_raw)  # Crear una versión JSON separada
            }
        }
        
        print(f"DEBUG - Contexto enviado al template: {contexto}")
        
        return render(request, 'simplex/simplex.html', contexto)
        
    except Exception as e:
        print(f"Error cargando historial: {e}")
        return redirect('simplex')

@login_required
def exportar_simplex_pdf(request, historial_id):
    """
    Genera un PDF con el historial completo de un cálculo del método Simplex.
    Solo disponible para usuarios premium.
    """
    # Verificar que el usuario sea premium
    if not (hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium):
        return HttpResponse("Acceso denegado. Esta funcionalidad requiere cuenta Premium.", status=403)
    
    try:
        # Obtener el historial del usuario actual
        historial = SimplexHistorial.objects.get(id=historial_id, user=request.user)
    except SimplexHistorial.DoesNotExist:
        return HttpResponse("Historial no encontrado.", status=404)
    
    # Función para limpiar texto LaTeX y convertirlo a texto legible
    def limpiar_latex(texto):
        if not texto:
            return ""
        
        # Diccionario de reemplazos LaTeX comunes
        import re
        replacements = {
            r'\\text\{([^}]+)\}': r'\1',  # \text{...} -> ...
            r'\\geq': '≥',
            r'\\leq': '≤',
            r'\\le': '≤',
            r'\\ge': '≥',
            r'\\cdot': '·',
            r'\\times': '×',
            r'\\div': '÷',
            r'\\pm': '±',
            r'\\mp': '∓',
            r'\\infty': '∞',
            r'\\alpha': 'α',
            r'\\beta': 'β',
            r'\\gamma': 'γ',
            r'\\delta': 'δ',
            r'\\epsilon': 'ε',
            r'\\theta': 'θ',
            r'\\lambda': 'λ',
            r'\\mu': 'μ',
            r'\\pi': 'π',
            r'\\sigma': 'σ',
            r'\\tau': 'τ',
            r'\\phi': 'φ',
            r'\\omega': 'ω',
            r'\\\{': '{',
            r'\\\}': '}',
            r'\\\\': '',
            r'\$': '',
            r'_\{([^}]+)\}': r'_\1',  # _{...} -> _...
            r'\^{([^}]+)}': r'^\1',   # ^{...} -> ^...
        }
        
        # Aplicar reemplazos
        for pattern, replacement in replacements.items():
            texto = re.sub(pattern, replacement, texto)
        
        # Limpiar espacios extra y caracteres problemáticos
        texto = re.sub(r'\s+', ' ', texto).strip()
        
        return texto
    
    # Crear respuesta PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="simplex_historial_{historial.id}.pdf"'
    
    # Crear el PDF
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    import json
    from datetime import datetime
    import re
    
    # Configurar documento
    doc = SimpleDocTemplate(response, pagesize=A4, 
                           rightMargin=72, leftMargin=72, 
                           topMargin=72, bottomMargin=18)
    
    # Estilos mejorados
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        textColor=colors.darkblue,
        alignment=1,  # Centrado
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=15,
        spaceBefore=20,
        textColor=colors.darkgreen,
        fontName='Helvetica-Bold'
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.darkred,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leftIndent=0
    )
    
    step_style = ParagraphStyle(
        'StepStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=5,
        leftIndent=20,
        bulletIndent=10
    )
    
    code_style = ParagraphStyle(
        'CodeStyle',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=5,
        leftIndent=30,
        fontName='Courier',
        textColor=colors.darkblue
    )
    
    # Contenido del PDF
    story = []
    
    # Título principal
    story.append(Paragraph("MÉTODO SIMPLEX - REPORTE COMPLETO", title_style))
    story.append(Paragraph("Programación Lineal - Algoritmo Simplex", normal_style))
    story.append(Spacer(1, 20))
    
    # Información básica
    story.append(Paragraph("INFORMACIÓN DEL CÁLCULO", heading_style))
    fecha_formato = historial.fecha.strftime("%d/%m/%Y %H:%M:%S")
    story.append(Paragraph(f"<b>Fecha de cálculo:</b> {fecha_formato}", normal_style))
    story.append(Paragraph(f"<b>ID del problema:</b> {historial.id}", normal_style))
    story.append(Paragraph(f"<b>Usuario:</b> {request.user.username}", normal_style))
    story.append(Spacer(1, 20))
    
    # Planteamiento del problema
    story.append(Paragraph("PLANTEAMIENTO DEL PROBLEMA", heading_style))
    
    # Función objetivo limpia
    funcion_limpia = limpiar_latex(historial.funcion_objetivo)
    story.append(Paragraph(f"<b>Función Objetivo:</b>", subheading_style))
    story.append(Paragraph(f"{historial.tipo_objetivo} Z = {funcion_limpia}", normal_style))
    story.append(Spacer(1, 10))
    
    # Restricciones limpias
    story.append(Paragraph(f"<b>Restricciones:</b>", subheading_style))
    try:
        restricciones = json.loads(historial.restricciones)
        for i, restriccion in enumerate(restricciones, 1):
            restriccion_limpia = limpiar_latex(restriccion)
            story.append(Paragraph(f"  {i}. {restriccion_limpia}", normal_style))
    except:
        restricciones_texto = limpiar_latex(historial.restricciones)
        story.append(Paragraph(f"  {restricciones_texto}", normal_style))
    
    # Restricciones de no negatividad - obtener nombres de variables reales
    try:
        solucion_temp = json.loads(historial.solucion_optima)
        variables_nombres = [var for var in solucion_temp.keys() if not (var.startswith('s') and var[1:].isdigit())]
        if variables_nombres:
            variables_texto = ", ".join(variables_nombres)
            story.append(Paragraph(f"  Restricciones de no negatividad: {variables_texto} ≥ 0", normal_style))
        else:
            story.append(Paragraph("  Restricciones de no negatividad: x₁, x₂, ... ≥ 0", normal_style))
    except:
        story.append(Paragraph("  Restricciones de no negatividad: x₁, x₂, ... ≥ 0", normal_style))
    story.append(Spacer(1, 20))
    
    # Modelo matemático completo (forma estándar)
    try:
        modelo = json.loads(historial.modelo_matematico)
        story.append(Paragraph("MODELO EN FORMA ESTÁNDAR", heading_style))
        
        if 'funcion_objetivo' in modelo:
            funcion_limpia_modelo = limpiar_latex(modelo['funcion_objetivo'])
            story.append(Paragraph(f"<b>Función objetivo:</b> {funcion_limpia_modelo}", normal_style))
        
        if 'funcion_z_estandar' in modelo:
            z_estandar_limpia = limpiar_latex(modelo['funcion_z_estandar'])
            story.append(Paragraph(f"<b>Ecuación Z:</b> {z_estandar_limpia}", normal_style))
        
        if 'restricciones' in modelo:
            story.append(Paragraph("<b>Sistema de ecuaciones con variables de holgura:</b>", subheading_style))
            for i, restriccion in enumerate(modelo['restricciones'], 1):
                restriccion_limpia = limpiar_latex(restriccion)
                story.append(Paragraph(f"  {i}. {restriccion_limpia}", normal_style))
        
        story.append(Spacer(1, 20))
    except:
        pass
    
    # Solución óptima detallada
    story.append(Paragraph("SOLUCIÓN ÓPTIMA", heading_style))
    
    try:
        solucion = json.loads(historial.solucion_optima)
        
        # Información general de la solución
        valor_optimo = historial.valor_z
        tipo_problema = 'máximo' if historial.tipo_objetivo == 'Maximizar' else 'mínimo'
        story.append(Paragraph(f"<b>Valor {tipo_problema} de Z:</b> {valor_optimo:.4f}", subheading_style))
        story.append(Spacer(1, 10))
        
        # Separar variables de decisión y de holgura
        variables_decision = {}
        variables_holgura = {}
        
        for var, valor in solucion.items():
            if var.startswith('s') and var[1:].isdigit():
                variables_holgura[var] = float(valor)
            else:
                variables_decision[var] = float(valor)
        
        # Tabla de variables de decisión
        if variables_decision:
            story.append(Paragraph("<b>Variables de Decisión:</b>", subheading_style))
            decision_data = [['Variable', 'Valor Óptimo', 'Interpretación']]
            for var, valor in variables_decision.items():
                interpretacion = f"Valor óptimo de {var}"
                decision_data.append([var, f"{valor:.4f}", interpretacion])
            
            decision_table = Table(decision_data, colWidths=[1.5*inch, 1.5*inch, 2.5*inch])
            decision_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(decision_table)
            story.append(Spacer(1, 15))
        
        # Tabla de variables de holgura
        if variables_holgura:
            story.append(Paragraph("<b>Variables de Holgura:</b>", subheading_style))
            holgura_data = [['Variable', 'Valor', 'Estado de la Restricción']]
            for var, valor in variables_holgura.items():
                if valor > 0.0001:
                    estado = f"Restricción no activa (sobra {valor:.4f})"
                else:
                    estado = "Restricción activa (en el límite)"
                holgura_data.append([var, f"{valor:.4f}", estado])
            
            holgura_table = Table(holgura_data, colWidths=[1.5*inch, 1.5*inch, 2.5*inch])
            holgura_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(holgura_table)
            story.append(Spacer(1, 20))
    
    except Exception as e:
        story.append(Paragraph(f"<b>Solución:</b> {historial.solucion_optima}", normal_style))
        story.append(Paragraph(f"<b>Valor de Z:</b> {historial.valor_z:.4f}", normal_style))
        story.append(Spacer(1, 20))
    
    # Proceso de resolución detallado
    try:
        iteraciones = json.loads(historial.iteraciones_json)
        if iteraciones:
            story.append(Paragraph("PROCESO DE RESOLUCIÓN - ALGORITMO SIMPLEX", heading_style))
            story.append(Paragraph(f"<b>Número total de iteraciones:</b> {len(iteraciones)}", normal_style))
            story.append(Paragraph("El algoritmo simplex encontró la solución óptima mediante las siguientes iteraciones:", normal_style))
            story.append(Spacer(1, 15))
            
            # Resumen de iteraciones
            for i, iteracion in enumerate(iteraciones):
                story.append(Paragraph(f"<b>Iteración {i}:</b>", subheading_style))
                
                if iteracion.get('es_optima', False):
                    story.append(Paragraph("• ✅ <b>Solución óptima encontrada</b>", step_style))
                    story.append(Paragraph("• Todos los coeficientes de la fila Z son no negativos", step_style))
                    story.append(Paragraph("• El algoritmo termina exitosamente", step_style))
                else:
                    if 'variable_entra' in iteracion and 'variable_sale' in iteracion:
                        var_entra = iteracion['variable_entra']
                        var_sale = iteracion['variable_sale']
                        elemento_pivote = iteracion.get('elemento_pivote', 'N/A')
                        
                        story.append(Paragraph(f"• Variable que entra a la base: <b>{var_entra}</b>", step_style))
                        story.append(Paragraph(f"• Variable que sale de la base: <b>{var_sale}</b>", step_style))
                        story.append(Paragraph(f"• Elemento pivote: <b>{elemento_pivote}</b>", step_style))
                        
                        if 'operaciones' in iteracion:
                            operaciones = iteracion['operaciones']
                            if operaciones:
                                story.append(Paragraph("• Operaciones realizadas:", step_style))
                                for op in operaciones[:3]:  # Limitar a las primeras 3 operaciones
                                    if op.get('tipo') == 'normalizacion':
                                        story.append(Paragraph(f"  - Normalización de fila {op.get('fila', '')}", code_style))
                
                story.append(Spacer(1, 10))
            
            story.append(Spacer(1, 20))
        else:
            story.append(Paragraph("PROCESO DE RESOLUCIÓN", heading_style))
            story.append(Paragraph("El problema fue resuelto usando el método simplex estándar.", normal_style))
            story.append(Paragraph("Las iteraciones detalladas están disponibles en la versión completa del sistema.", normal_style))
            story.append(Spacer(1, 20))
    except:
        story.append(Paragraph("PROCESO DE RESOLUCIÓN", heading_style))
        story.append(Paragraph("El problema fue resuelto exitosamente usando el algoritmo simplex.", normal_style))
        story.append(Spacer(1, 20))
    
    # Interpretación de resultados
    story.append(Paragraph("INTERPRETACIÓN DE RESULTADOS", heading_style))
    
    try:
        solucion = json.loads(historial.solucion_optima)
        variables_decision = {k: v for k, v in solucion.items() if not (k.startswith('s') and k[1:].isdigit())}
        
        if variables_decision:
            story.append(Paragraph(f"La solución óptima del problema de {historial.tipo_objetivo.lower()} indica que:", normal_style))
            story.append(Spacer(1, 10))
            
            for var, valor in variables_decision.items():
                story.append(Paragraph(f"• La variable <b>{var}</b> debe tomar el valor <b>{valor:.4f}</b>", step_style))
            
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"Con esta asignación de variables, el valor {'máximo' if historial.tipo_objetivo == 'Maximizar' else 'mínimo'} de la función objetivo Z es <b>{historial.valor_z:.4f}</b>.", normal_style))
            
            # Análisis de holguras
            variables_holgura = {k: v for k, v in solucion.items() if k.startswith('s') and k[1:].isdigit()}
            if variables_holgura:
                story.append(Spacer(1, 15))
                story.append(Paragraph("<b>Análisis de restricciones:</b>", subheading_style))
                
                restricciones_activas = 0
                for var, valor in variables_holgura.items():
                    num_restriccion = var[1:]
                    if float(valor) < 0.0001:
                        story.append(Paragraph(f"• Restricción {num_restriccion}: <b>Activa</b> (se cumple con igualdad)", step_style))
                        restricciones_activas += 1
                    else:
                        story.append(Paragraph(f"• Restricción {num_restriccion}: No activa (sobran {valor:.4f} unidades)", step_style))
                
                story.append(Spacer(1, 10))
                story.append(Paragraph(f"Total de restricciones activas: <b>{restricciones_activas}</b>", normal_style))
    
    except:
        story.append(Paragraph(f"El problema de {historial.tipo_objetivo.lower()} fue resuelto exitosamente.", normal_style))
        story.append(Paragraph(f"El valor óptimo de la función objetivo es: <b>{historial.valor_z:.4f}</b>", normal_style))
    
    story.append(Spacer(1, 20))
    
    # Información técnica
    story.append(Paragraph("INFORMACIÓN TÉCNICA", heading_style))
    story.append(Paragraph("<b>Método utilizado:</b> Algoritmo Simplex Estándar", normal_style))
    story.append(Paragraph("<b>Tipo de problema:</b> Programación Lineal", normal_style))
    story.append(Paragraph("<b>Forma estándar:</b> Problema convertido con variables de holgura", normal_style))
    story.append(Paragraph("<b>Precisión numérica:</b> Cálculos con precisión de punto flotante", normal_style))
    
    try:
        solucion = json.loads(historial.solucion_optima)
        variables_decision = len([k for k in solucion.keys() if not (k.startswith('s') and k[1:].isdigit())])
        variables_holgura = len([k for k in solucion.keys() if k.startswith('s') and k[1:].isdigit()])
        story.append(Paragraph(f"<b>Variables de decisión:</b> {variables_decision}", normal_style))
        story.append(Paragraph(f"<b>Variables de holgura:</b> {variables_holgura}", normal_style))
    except:
        pass
    
    story.append(Spacer(1, 30))
    
    # Pie de página mejorado
    story.append(Paragraph("_" * 80, normal_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("FORJA SOLVER", ParagraphStyle('FooterTitle', parent=styles['Normal'], fontSize=12, textColor=colors.darkblue, fontName='Helvetica-Bold', alignment=1)))
    story.append(Paragraph("Sistema de Resolución de Métodos Numéricos", ParagraphStyle('FooterSubtitle', parent=styles['Normal'], fontSize=10, textColor=colors.grey, alignment=1)))
    story.append(Spacer(1, 5))
    story.append(Paragraph(f"Reporte generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", ParagraphStyle('FooterDate', parent=styles['Normal'], fontSize=9, textColor=colors.grey, alignment=1)))
    
    # Construir PDF
    doc.build(story)
    
    return response

@login_required
def metodo_grafico_pdf(request, id):
    """
    Genera un PDF con el historial completo de un cálculo del método gráfico.
    Solo disponible para usuarios premium.
    """
    # Verificar que el usuario sea premium
    if not (hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium):
        return HttpResponse("Acceso denegado. Esta funcionalidad requiere cuenta Premium.", status=403)
    
    try:
        # Obtener el historial del usuario actual
        historial = MetodoGraficoHistorial.objects.get(id=id, user=request.user)
    except MetodoGraficoHistorial.DoesNotExist:
        return HttpResponse("Historial no encontrado.", status=404)
    
    # Función para limpiar texto LaTeX y convertirlo a texto legible
    def limpiar_latex(texto):
        if not texto:
            return ""
        
        # Diccionario de reemplazos LaTeX comunes
        replacements = {
            r'\\text\{([^}]+)\}': r'\1',  # \text{...} -> ...
            r'\\geq': '≥',
            r'\\leq': '≤',
            r'\\le': '≤',
            r'\\ge': '≥',
            r'\\cdot': '·',
            r'\\times': '×',
            r'\\div': '÷',
            r'\\pm': '±',
            r'\\mp': '∓',
            r'\\infty': '∞',
            r'\\alpha': 'α',
            r'\\beta': 'β',
            r'\\gamma': 'γ',
            r'\\delta': 'δ',
            r'\\epsilon': 'ε',
            r'\\theta': 'θ',
            r'\\lambda': 'λ',
            r'\\mu': 'μ',
            r'\\pi': 'π',
            r'\\sigma': 'σ',
            r'\\tau': 'τ',
            r'\\phi': 'φ',
            r'\\omega': 'ω',
            r'\\\{': '{',
            r'\\\}': '}',
            r'\\\\': '',
            r'\$': '',
            r'_\{([^}]+)\}': r'_\1',  # _{...} -> _...
            r'\^{([^}]+)}': r'^\1',   # ^{...} -> ^...
        }
        
        # Aplicar reemplazos
        import re
        for pattern, replacement in replacements.items():
            texto = re.sub(pattern, replacement, texto)
        
        # Limpiar espacios extra y caracteres problemáticos
        texto = re.sub(r'\s+', ' ', texto).strip()
        
        return texto
    
    # Crear respuesta PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="metodo_grafico_{historial.id}.pdf"'
    
    # Crear el PDF
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    import json
    from datetime import datetime
    import re
    
    # Configurar documento
    doc = SimpleDocTemplate(response, pagesize=A4, 
                           rightMargin=72, leftMargin=72, 
                           topMargin=72, bottomMargin=18)
    
    # Estilos mejorados
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        textColor=colors.darkblue,
        alignment=1,  # Centrado
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=15,
        spaceBefore=20,
        textColor=colors.darkgreen,
        fontName='Helvetica-Bold'
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.darkred,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        leftIndent=0
    )
    
    step_style = ParagraphStyle(
        'StepStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=5,
        leftIndent=20,
        bulletIndent=10
    )
    
    # Contenido del PDF
    story = []
    
    # Título principal
    story.append(Paragraph("MÉTODO GRÁFICO - REPORTE COMPLETO", title_style))
    story.append(Paragraph("Programación Lineal - Resolución por Método Gráfico", normal_style))
    story.append(Spacer(1, 20))
    
    # Información básica
    story.append(Paragraph("INFORMACIÓN DEL CÁLCULO", heading_style))
    fecha_formato = historial.fecha_creacion.strftime("%d/%m/%Y %H:%M:%S")
    story.append(Paragraph(f"<b>Fecha de cálculo:</b> {fecha_formato}", normal_style))
    story.append(Paragraph(f"<b>ID del problema:</b> {historial.id}", normal_style))
    story.append(Paragraph(f"<b>Usuario:</b> {request.user.username}", normal_style))
    story.append(Spacer(1, 20))
    
    # Planteamiento del problema
    story.append(Paragraph("PLANTEAMIENTO DEL PROBLEMA", heading_style))
    
    # Función objetivo limpia
    funcion_limpia = limpiar_latex(historial.funcion)
    tipo_optimizacion = 'Maximizar' if historial.optimizacion == 'max' else 'Minimizar'
    story.append(Paragraph(f"<b>Función Objetivo:</b>", subheading_style))
    story.append(Paragraph(f"{tipo_optimizacion} Z = {funcion_limpia}", normal_style))
    story.append(Spacer(1, 10))
    
    # Restricciones limpias
    story.append(Paragraph(f"<b>Restricciones:</b>", subheading_style))
    restricciones_lineas = historial.restricciones.strip().split('\n')
    for i, restriccion in enumerate(restricciones_lineas, 1):
        if restriccion.strip():
            restriccion_limpia = limpiar_latex(restriccion.strip())
            story.append(Paragraph(f"  {i}. {restriccion_limpia}", normal_style))
    
    # Restricciones de no negatividad - obtener nombres de variables reales
    try:
        # Para método gráfico normalmente son 2 variables
        if hasattr(historial, 'vertices') and historial.vertices:
            # Intentar extraer nombres de variables de la función objetivo
            import re
            funcion_texto = historial.funcion if hasattr(historial, 'funcion') else ""
            variables_match = re.findall(r'([a-zA-Z_][a-zA-Z0-9_]*)', funcion_texto)
            variables_unicas = []
            for var in variables_match:
                if var not in ['text', 'max', 'min'] and var not in variables_unicas:
                    variables_unicas.append(var)
            
            if len(variables_unicas) >= 2:
                variables_texto = ", ".join(variables_unicas[:2])  # Solo las primeras 2
                story.append(Paragraph(f"  Restricciones de no negatividad: {variables_texto} ≥ 0", normal_style))
            else:
                story.append(Paragraph("  Restricciones de no negatividad: x ≥ 0, y ≥ 0", normal_style))
        else:
            story.append(Paragraph("  Restricciones de no negatividad: x ≥ 0, y ≥ 0", normal_style))
    except:
        story.append(Paragraph("  Restricciones de no negatividad: x ≥ 0, y ≥ 0", normal_style))
    story.append(Spacer(1, 20))
    
    # Análisis de la región factible
    story.append(Paragraph("ANÁLISIS DE LA REGIÓN FACTIBLE", heading_style))
    
    # Vértices encontrados
    if historial.vertices:
        story.append(Paragraph("<b>Vértices de la región factible:</b>", subheading_style))
        vertices_lineas = historial.vertices.strip().split('\n')
        for i, vertice in enumerate(vertices_lineas, 1):
            if vertice.strip():
                story.append(Paragraph(f"  V{i}: {vertice.strip()}", normal_style))
        story.append(Spacer(1, 15))
    
    # Validación de puntos factibles
    if historial.puntos_factibles:
        story.append(Paragraph("<b>Evaluación de puntos factibles:</b>", subheading_style))
        puntos_lineas = historial.puntos_factibles.strip().split('\n')
        for punto in puntos_lineas:
            if punto.strip():
                story.append(Paragraph(f"  • {punto.strip()}", normal_style))
        story.append(Spacer(1, 15))
    
    # Procedimiento paso a paso
    if historial.iteraciones:
        story.append(Paragraph("PROCEDIMIENTO DETALLADO PASO A PASO", heading_style))
        
        # Procesar el texto de iteraciones para hacerlo más legible
        procedimiento_texto = historial.iteraciones.strip()
        
        # Dividir en secciones por pasos numerados
        secciones = procedimiento_texto.split('\n\n')
        
        for seccion in secciones:
            if seccion.strip():
                lineas = seccion.strip().split('\n')
                
                # Identificar si es un título de paso
                primera_linea = lineas[0].strip()
                if re.match(r'^\d+\.', primera_linea):
                    # Es un paso numerado - usar como subtítulo
                    titulo_limpio = limpiar_latex(primera_linea)
                    story.append(Paragraph(titulo_limpio, subheading_style))
                    
                    # Procesar el resto de líneas del paso
                    for linea in lineas[1:]:
                        if linea.strip():
                            linea_limpia = limpiar_latex(linea.strip())
                            if linea_limpia.startswith('•') or linea_limpia.startswith('-'):
                                story.append(Paragraph(f"  {linea_limpia}", step_style))
                            else:
                                story.append(Paragraph(linea_limpia, normal_style))
                else:
                    # No es un paso numerado, procesar todas las líneas
                    for linea in lineas:
                        if linea.strip():
                            linea_limpia = limpiar_latex(linea.strip())
                            if ':' in linea_limpia and not linea_limpia.startswith(' '):
                                # Parece ser un subtítulo
                                story.append(Paragraph(f"<b>{linea_limpia}</b>", normal_style))
                            else:
                                story.append(Paragraph(linea_limpia, normal_style))
                
                story.append(Spacer(1, 10))
        
        story.append(Spacer(1, 20))
    
    # Solución óptima destacada
    story.append(Paragraph("SOLUCIÓN ÓPTIMA", heading_style))
    
    # Crear tabla con la solución
    punto_limpio = limpiar_latex(historial.punto_optimo)
    solucion_limpia = limpiar_latex(historial.solucion)
    
    # Datos de la tabla de solución
    solution_data = [
        ['Concepto', 'Valor'],
        ['Punto Óptimo (x, y)', punto_limpio],
        ['Valor de la Función Objetivo', solucion_limpia.replace('Z = ', '')],
        ['Tipo de Optimización', tipo_optimizacion]
    ]
    
    solution_table = Table(solution_data, colWidths=[3*inch, 2*inch])
    solution_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    story.append(solution_table)
    story.append(Spacer(1, 20))
    
    # Interpretación de resultados
    story.append(Paragraph("INTERPRETACIÓN DE RESULTADOS", heading_style))
    story.append(Paragraph(f"El problema de {tipo_optimizacion.lower()} la función objetivo Z = {funcion_limpia} sujeto a las restricciones dadas tiene como solución óptima:", normal_style))
    story.append(Spacer(1, 10))
    
    # Extraer coordenadas del punto óptimo para interpretación
    try:
        punto_coords = re.findall(r'[-+]?\d*\.?\d+', punto_limpio)
        if len(punto_coords) >= 2:
            x_opt = float(punto_coords[0])
            y_opt = float(punto_coords[1])
            story.append(Paragraph(f"• <b>Variable x:</b> {x_opt:.4f}", normal_style))
            story.append(Paragraph(f"• <b>Variable y:</b> {y_opt:.4f}", normal_style))
    except:
        story.append(Paragraph(f"• <b>Punto óptimo:</b> {punto_limpio}", normal_style))
    
    # Extraer valor óptimo
    try:
        valor_z = re.findall(r'[-+]?\d*\.?\d+', solucion_limpia)
        if valor_z:
            z_opt = float(valor_z[0])
            story.append(Paragraph(f"• <b>Valor {'máximo' if historial.optimizacion == 'max' else 'mínimo'} de Z:</b> {z_opt:.4f}", normal_style))
    except:
        story.append(Paragraph(f"• <b>Valor óptimo:</b> {solucion_limpia}", normal_style))
    
    story.append(Spacer(1, 15))
    story.append(Paragraph("Este resultado representa la mejor solución posible dentro de la región factible definida por las restricciones del problema.", normal_style))
    story.append(Spacer(1, 20))
    
    # Nota sobre la gráfica
    if historial.grafica:
        story.append(Paragraph("REPRESENTACIÓN GRÁFICA", heading_style))
        story.append(Paragraph("La representación gráfica interactiva de la región factible, las restricciones y la solución óptima está disponible en la versión digital del sistema.", normal_style))
        story.append(Paragraph("La gráfica muestra:", normal_style))
        story.append(Paragraph("• Las líneas de restricción que delimitan la región factible", step_style))
        story.append(Paragraph("• La región factible sombreada", step_style))
        story.append(Paragraph("• Los vértices de la región factible marcados como puntos", step_style))
        story.append(Paragraph("• El punto óptimo destacado con una estrella dorada", step_style))
        story.append(Spacer(1, 20))
    
    # Información adicional
    story.append(Paragraph("INFORMACIÓN TÉCNICA", heading_style))
    story.append(Paragraph("<b>Método utilizado:</b> Método Gráfico para Programación Lineal", normal_style))
    story.append(Paragraph("<b>Algoritmo:</b> Evaluación exhaustiva de vértices de la región factible", normal_style))
    story.append(Paragraph("<b>Precisión:</b> Cálculos realizados con precisión de punto flotante", normal_style))
    story.append(Paragraph(f"<b>Número de restricciones:</b> {len(restricciones_lineas) + 2} (incluyendo no negatividad)", normal_style))
    story.append(Paragraph(f"<b>Variables de decisión:</b> 2 (x, y)", normal_style))
    story.append(Spacer(1, 30))
    
    # Pie de página mejorado
    story.append(Paragraph("_" * 80, normal_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("FORJA SOLVER", ParagraphStyle('FooterTitle', parent=styles['Normal'], fontSize=12, textColor=colors.darkblue, fontName='Helvetica-Bold', alignment=1)))
    story.append(Paragraph("Sistema de Resolución de Métodos Numéricos", ParagraphStyle('FooterSubtitle', parent=styles['Normal'], fontSize=10, textColor=colors.grey, alignment=1)))
    story.append(Spacer(1, 5))
    story.append(Paragraph(f"Reporte generado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", ParagraphStyle('FooterDate', parent=styles['Normal'], fontSize=9, textColor=colors.grey, alignment=1)))
    
    # Construir PDF
    doc.build(story)
    
    return response

@login_required
def repetir_metodo_grafico(request, id):
    """
    Cargar un problema del historial del método gráfico para resolverlo nuevamente
    """
    try:
        historial_item = get_object_or_404(MetodoGraficoHistorial, id=id, user=request.user)
        
        # Preparar datos para precargar el formulario
        contexto = {
            'historial_data': {
                'objective': historial_item.funcion,
                'optimization': historial_item.optimizacion,
                'restrictions': historial_item.restricciones
            },
            'default_function': historial_item.funcion,
            'default_restrictions': historial_item.restricciones,
            'default_optimization': historial_item.optimizacion,
            'steps': [],
            'errors': [],
            'show_results': False,
            'is_premium': hasattr(request.user, 'userprofile') and request.user.userprofile.is_premium,
            'show_graph': False
        }
        
        return render(request, 'metodo_grafico/metodo_grafico.html', contexto)
        
    except Exception as e:
        print(f"Error cargando historial del método gráfico: {e}")
        return redirect('metodoGrafico')

#### gran m 


from django.shortcuts import render
from .mop import GranMSimplexExtended
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.utils.html import strip_tags
from bs4 import BeautifulSoup
from django.utils.html import strip_tags
from openpyxl import Workbook
from .models import GranMHistorial
import json

@login_required
@csrf_exempt
def metodo_gran_m(request):
    resultado_html = None

    if request.method == 'POST':
        tipo = request.POST.get('tipo', 'min')
        objetivo = request.POST.get('objetivo', '')
        minimizar = (tipo == 'min')

        try:
            obj_raw = objetivo.split(',')
            if len(obj_raw) > 10:
                raise ValueError("Máximo 10 variables permitidas en la función objetivo.")

            obj = [float(v.strip()) for v in obj_raw if v.strip() != '']

            restricciones = []
            tipos = []

            for i in range(10):
                restr = request.POST.get(f'restriccion_{i}')
                signo = request.POST.get(f'signo_{i}')
                if restr and signo:
                    partes = [x.strip() for x in restr.split(',')]
                    if len(partes) != len(obj) + 1:
                        raise ValueError(f"La restricción {i+1} debe tener {len(obj)+1} valores.")
                    valores = [float(x) for x in partes]
                    restricciones.append(valores)
                    tipos.append(signo)

            if len(restricciones) == 0:
                raise ValueError("Debes ingresar al menos una restricción.")

            solver = GranMSimplexExtended()
            resultado_html = solver.solve(obj, restricciones, tipos, minimizar)

            soup = BeautifulSoup(resultado_html, "html.parser")
            for s in soup(["style", "script"]):
                s.decompose()
            request.session['resultado_exportable'] = str(soup)

            # Guardar historial
            GranMHistorial.objects.create(
                user=request.user,
                tipo=tipo,
                funcion_objetivo=objetivo,
                restricciones=json.dumps(restricciones),
                signos=json.dumps(tipos),
                resultado_html=str(soup)
            )

        except ValueError as ve:
            resultado_html = f"<p style='color:red;'>Error de validación: {str(ve)}</p>"
        except Exception as e:
            resultado_html = f"<p style='color:red;'>Error inesperado: {str(e)}</p>"

    return render(request, 'paginas/metodoGranM.html', {'resultado': resultado_html})

from .models import GranMHistorial
from django.contrib.auth.decorators import login_required

@login_required
def historial_gran_m(request):
    historial = GranMHistorial.objects.filter(user=request.user).order_by('-fecha')
    return render(request, 'paginas/historialGranM.html', {'historial': historial})

# Exportar a PDF
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from django.template.loader import render_to_string
from xhtml2pdf import pisa
import io

def exportar_pdf(request):
    resultado = request.session.get('resultado_exportable', '')
    if not resultado:
        return HttpResponse("No hay resultado para exportar.", status=400)

    html = render_to_string("paginas/export_template.html", {'resultado': resultado})
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="resultado.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Error al generar el PDF', status=500)
    return response


# Exportar a Excel
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

def exportar_excel(request):
    html = request.session.get('resultado_exportable', '')
    soup = BeautifulSoup(html, "html.parser")

    wb = Workbook()
    ws = wb.active
    ws.title = "Gran M Resultado"
    row = 1

    bold_font = Font(bold=True)
    blue_font = Font(bold=True, color="1F4E78")
    gray_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    wrap_center = Alignment(wrap_text=True, vertical="top")

    encabezado = soup.find('h1') or soup.find('h2')
    if encabezado:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        cell = ws.cell(row=row, column=1, value=encabezado.get_text(strip=True))
        cell.font = Font(size=14, bold=True)
        row += 2

    for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table']):
        if tag.name in ['h1', 'h2']:
            cell = ws.cell(row=row, column=1, value=tag.get_text(strip=True))
            cell.font = blue_font
            row += 1
        elif tag.name == 'h3':
            cell = ws.cell(row=row, column=1, value=tag.get_text(strip=True))
            cell.font = Font(bold=True, color="4BACC6")
            row += 1
        elif tag.name == 'p':
            text = tag.get_text(strip=True)
            if text:
                cell = ws.cell(row=row, column=1, value=text)
                cell.alignment = wrap_center
                row += 1
        elif tag.name in ['ul', 'ol']:
            for li in tag.find_all('li'):
                cell = ws.cell(row=row, column=1, value=f"• {li.get_text(strip=True)}")
                cell.alignment = wrap_center
                row += 1
        elif tag.name == 'table':
            for tr in tag.find_all('tr'):
                col = 1
                for td in tr.find_all(['th', 'td']):
                    text = td.get_text(strip=True)
                    cell = ws.cell(row=row, column=col, value=text)
                    cell.alignment = wrap_center
                    if td.name == 'th':
                        cell.font = bold_font
                        cell.fill = gray_fill
                    col += 1
                row += 1
            row += 1

    for i, column_cells in enumerate(ws.columns, start=1):
        max_length = 0
        for cell in column_cells:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                continue
        col_letter = get_column_letter(i)
        ws.column_dimensions[col_letter].width = max(20, min(max_length + 5, 60))

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="resultado_estilizado.xlsx"'}
    )


# Exportar a Word
from docx import Document
from docx.shared import Pt

def exportar_word(request):
    html = request.session.get('resultado_exportable', '')
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    encabezado = soup.find('h2') or soup.find('h1')
    if encabezado:
        titulo = doc.add_heading(encabezado.get_text(), level=1)

    for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'table']):
        if tag.name in ['h1', 'h2', 'h3']:
            doc.add_heading(tag.get_text(strip=True), level=2)
        elif tag.name == 'p':
            doc.add_paragraph(tag.get_text(strip=True))
        elif tag.name in ['ul', 'ol']:
            for li in tag.find_all('li'):
                doc.add_paragraph(f"• {li.get_text(strip=True)}", style='ListBullet')
        elif tag.name == 'table':
            rows = tag.find_all('tr')
            if rows:
                cols = rows[0].find_all(['td', 'th'])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                for i, tr in enumerate(rows):
                    for j, td in enumerate(tr.find_all(['td', 'th'])):
                        table.cell(i, j).text = td.get_text(strip=True)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return HttpResponse(output,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename="resultado.docx"'}
    )


from django.utils.html import escape
import json

@login_required
@csrf_exempt
def editar_gran_m(request, pk):
    registro = get_object_or_404(GranMHistorial, pk=pk, user=request.user)
    resultado_html = None

    if request.method == 'POST':
        tipo = request.POST.get('tipo', 'min')
        objetivo = request.POST.get('objetivo', '')
        minimizar = (tipo == 'min')

        try:
            obj_raw = objetivo.split(',')
            if len(obj_raw) > 10:
                raise ValueError("Máximo 10 variables permitidas.")

            obj = [float(v.strip()) for v in obj_raw if v.strip()]
            restricciones = []
            tipos = []

            for i in range(10):
                restr = request.POST.get(f'restriccion_{i}')
                signo = request.POST.get(f'signo_{i}')
                if restr and signo:
                    partes = [x.strip() for x in restr.split(',')]
                    if len(partes) != len(obj) + 1:
                        raise ValueError(f"Restricción {i+1} inválida.")
                    restricciones.append([float(x) for x in partes])
                    tipos.append(signo)

            if not restricciones:
                raise ValueError("Al menos una restricción es requerida.")

            solver = GranMSimplexExtended()
            resultado_html = solver.solve(obj, restricciones, tipos, minimizar)

            GranMHistorial.objects.create(
                user=request.user,
                tipo=tipo,
                funcion_objetivo=objetivo,
                restricciones=json.dumps(restricciones),  
                signos=json.dumps(tipos),                
                resultado_html=resultado_html
            )

            soup = BeautifulSoup(resultado_html, "html.parser")
            for tag in soup(["style", "script"]): tag.decompose()
            request.session['resultado_exportable'] = str(soup)

        except Exception as e:
            resultado_html = f"<p style='color:red;'>Error: {str(e)}</p>"

    try:
        restricciones_parseadas = json.loads(registro.restricciones)
        signos_parseados = json.loads(registro.signos)
    except json.JSONDecodeError as e:
        restricciones_parseadas = []
        signos_parseados = []
        resultado_html = f"<p style='color:red;'>Error cargando historial: {str(e)}</p>"

    return render(request, 'paginas/metodoGranM.html', {
        'resultado': resultado_html,
        'editar': True,
        'registro': {
            'tipo': registro.tipo,
            'funcion_objetivo': registro.funcion_objetivo,
            'restricciones': restricciones_parseadas,
            'signos': signos_parseados
        }
    })
